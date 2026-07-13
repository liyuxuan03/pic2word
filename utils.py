import os
import re
import math
import glob
import copy
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import fitz
import numpy as np

try:
    import yaml
except ImportError as e:
    raise ImportError("缺少 PyYAML 依赖，请先安装：pip install pyyaml") from e


# 高级版完整 schema 模板：当用户未提供 report_schema.yaml 时，quickstart 模式将以此为基础展开。
FULL_SCHEMA_TEMPLATE = {'schema_version': '2.0',
 'project': {'name': 'MintPy Word Auto Report', 'language': 'zh-CN', 'description': '基于 MintPy 结果自动生成 Word 报告'},
 'runtime_paths': {'logo_path': '/home/liyuxuan/penguin/pic2word/pic/logo.png',
                   'base_dir': '/home/liyuxuan/penguin/data/Galapagos_s1/mintpy',
                   'pic_dir': '{runtime_paths.base_dir}/pic',
                   'output_dir': '/home/liyuxuan/penguin/pic2word/vision_2/word',
                   'velocity_file': '{runtime_paths.base_dir}/velocity.h5',
                   'cfg_file': '{runtime_paths.pic_dir}/smallbaselineApp.cfg',
                   'rms_txt': '{runtime_paths.pic_dir}/rms_timeseriesResidual_ramp.txt',
                   'reference_date_txt': '{runtime_paths.pic_dir}/reference_date.txt',
                   'exclude_date_txt': '{runtime_paths.pic_dir}/exclude_date.txt'},
 'defaults': {'font_cn': '仿宋',
              'font_en': 'Times New Roman',
              'figure_width_cm': 14.0,
              'figure_max_height_cm': 14.0,
              'table_style': 'Table Grid',
              'missing_asset_policy': 'warn',
              'auto_page_break_after_chapter': True,
              'pdf_render_dpi': 300,
              'figure_numbering_scope': 'chapter',
              'table_numbering_scope': 'chapter'},
 'data_sources': {'mintpy_meta': {'type': 'mintpy_attribute', 'path': '{runtime_paths.velocity_file}'},
                  'obs_dates': {'type': 'text_file', 'path': '{runtime_paths.rms_txt}', 'parser': 'date_list'},
                  'mintpy_cfg': {'type': 'cfg_file', 'path': '{runtime_paths.cfg_file}'},
                  'reference_dates': {'type': 'text_file',
                                      'path': '{runtime_paths.reference_date_txt}',
                                      'parser': 'inline_text'},
                  'exclude_dates': {'type': 'text_file',
                                    'path': '{runtime_paths.exclude_date_txt}',
                                    'parser': 'inline_text'},
                  'ts_method': {'type': 'interactive', 'prompt': '>>> 请输入时序估计方法 (默认: SBAS): ', 'default': 'SBAS'},
                  'analysis_text': {'type': 'interactive', 'prompt': '>>> 请在此处输入分析内容: ', 'default': '（请补充解译分析内容）'},
                  'poi_figures': {'type': 'glob',
                                  'pattern': '{runtime_paths.pic_dir}/final_figure*.png',
                                  'sort': 'natural'}},
 'assets': {'network_figure': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/network.pdf'},
            'coherence_history': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/coherenceHistory.pdf'},
            'coherence_matrix': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/coherenceMatrix.pdf'},
            'avg_spatial_coh': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/geo_avgSpatialCoh.png'},
            'temporal_coh': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/geo_temporalCoherence.png'},
            'mask_conn_comp': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/maskConnComp.png'},
            'avg_phase_velocity': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/avgPhaseVelocity.png'},
            'unwrap_error': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/numTriNonzeroIntAmbiguity.png'},
            'geometry_radar': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/geometryRadar.png'},
            'residual_rms': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/rms_timeseriesResidual_ramp.pdf'},
            'velocity_map': {'type': 'figure',
                             'candidates': ['{runtime_paths.pic_dir}/velocity_DEM.png',
                                            '{runtime_paths.pic_dir}/velocity.png']},
            'velocity_std': {'type': 'figure', 'path': '{runtime_paths.pic_dir}/velocity_std_plot.png'},
            'poi_series': {'type': 'figure_group', 'source': 'poi_figures'}},
 'templates': {'headings': {'cover_title': '基于{user_inputs.data_source}数据的{user_inputs.region_name}地区\n'
                                           '时序InSAR地表形变处理与分析报告'},
               'captions': {'network': '图{figure_no} '
                                       '干涉图网络。图中节点表示SAR影像的获取日期，连线表示用于生成干涉图的影像对。横轴为时间，纵轴为相对于参考影像的垂直基线（单位：米）。连线的颜色映射表示各干涉对的平均空间相干性数值，取值范围为0.2至1.0。',
                            'coherence_history': '图{figure_no} '
                                                 '相干性极值统计。图中横轴表示SAR影像的获取日期，纵轴表示平均空间相干性系数，取值范围为0至1。对于每一个观测日期，柱状图展示了与该影像相关联的所有干涉对的相干性极值统计。',
                            'coherence_matrix': '图{figure_no} 相干性矩阵。图中横轴与纵轴均表示SAR影像按获取时间先后排序的索引编号。',
                            'avg_spatial_coh': '图{figure_no} 地理编码后的平均空间相干性图。',
                            'temporal_coh': '图{figure_no} 地理编码后时间相干性图。',
                            'mask_conn_comp': '图{figure_no} 相位解缠连通分量观测掩膜。',
                            'avg_phase_velocity': '图{figure_no} 平均相位堆栈速率图。',
                            'unwrap_error': '图{figure_no} 相位解缠误差评估图。',
                            'geometry_radar': '图{figure_no} 雷达几何参数与辅助数据图。',
                            'residual_rms': '图{figure_no} 残余相位 RMS 可视化结果。',
                            'velocity_map': '图{figure_no} 形变速率图',
                            'velocity_std': '图{figure_no} VelocityStd',
                            'poi_timeseries': '图{figure_no} 感兴趣目标点{item_index}的时序形变图'},
               'text': {'reference_dates': '参考日期：{value}', 'exclude_dates': '剔除日期：{value}'}},
 'figure_generation': {
     'mode': 'auto',
     'manual_points': [],
     'auto_params': {
         'temporal_coh_threshold': 0.70,
         'min_region_pixels': 30,
         'max_regions': 10,
         'region_percentiles': [99.5, 99.0, 98.5, 98.0, 97.0, 95.0],
     },
     'plot': {
         'sub_lat': [36.21, 36.71],
         'sub_lon': [112.99, 114.05],
         'wrap_range': [-3, 3],
         'shade_az': 315,
         'shade_alt': 45,
         'shade_frac': 0.8,
         'base_color': 0.5,
         'shade_exag': 0.3,
         'scalebar': [0.2, 0.8, 0.1],
     },
 },
 'report_structure': [{'type': 'cover',
                       'params': {'data_source': 'interactive',
                                  'region_name': 'interactive',
                                  'lead_author': 'interactive',
                                  'logo_path': '{runtime_paths.logo_path}',
                                  'output_name_template': '{user_inputs.region_name}InSAR处理报告.docx'},
                       'default_data_source': 'Sentinel-1',
                       'default_region_name': '北京',
                       'default_lead_author': 'yourname'},
                      {'type': 'toc'},
                      {'type': 'chapter',
                       'id': 'chapter_1',
                       'title': '1. 数据情况',
                       'children': [{'type': 'metadata_table',
                                     'source': 'mintpy_meta',
                                     'fields': [['卫 星', 'satellite'],
                                                ['波 段', 'band'],
                                                ['波 长', 'wavelength'],
                                                ['成像模式', 'mode'],
                                                ['swath_Num', 'swath_Num'],
                                                ['空间分辨率', 'resolution'],
                                                ['相对轨道号', 'track'],
                                                ['升降轨', 'direction'],
                                                ['中心入射角', 'inc_angle'],
                                                ['中心方位角', 'az_angle'],
                                                ['极化方式', 'pol'],
                                                ['开始时间', 'start_time'],
                                                ['结束时间', 'end_time'],
                                                ['干涉对数量', 'num_ifgram'],
                                                ['经纬度范围', 'bbox']]},
                                    {'type': 'table',
                                     'title': '表{table_no} 选用SAR数据日期',
                                     'source': 'obs_dates',
                                     'renderer': 'date_grid',
                                     'columns': 4}]},
                      {'type': 'chapter',
                       'id': 'chapter_2',
                       'title': '2. 处理参数设置',
                       'children': [{'type': 'software_versions',
                                     'label': '处理软件',
                                     'platform_source': 'mintpy_meta',
                                     'stack_processor': 'auto'},
                                    {'type': 'kv_line', 'label': '时序估计方法', 'source': 'ts_method'},
                                    {'type': 'cfg_table',
                                     'title': 'MintPy 参数配置',
                                     'source': 'mintpy_cfg',
                                     'filter': {'exclude_values': ['auto', 'no']}}]},
                      {'type': 'chapter',
                       'id': 'chapter_3',
                       'title': '3. 干涉数据质量评估',
                       'children': [{'type': 'section',
                                     'id': 'chapter_3_1',
                                     'title': '3.1 干涉图网络',
                                     'children': [{'type': 'figure',
                                                   'asset': 'network_figure',
                                                   'caption': '{templates.captions.network}'},
                                                  {'type': 'figure',
                                                   'asset': 'coherence_history',
                                                   'caption': '{templates.captions.coherence_history}'},
                                                  {'type': 'figure',
                                                   'asset': 'coherence_matrix',
                                                   'caption': '{templates.captions.coherence_matrix}'}]},
                                    {'type': 'section',
                                     'id': 'chapter_3_2',
                                     'title': '3.2 相干性',
                                     'children': [{'type': 'figure',
                                                   'asset': 'avg_spatial_coh',
                                                   'caption': '{templates.captions.avg_spatial_coh}'},
                                                  {'type': 'figure',
                                                   'asset': 'temporal_coh',
                                                   'caption': '{templates.captions.temporal_coh}'},
                                                  {'type': 'figure',
                                                   'asset': 'mask_conn_comp',
                                                   'caption': '{templates.captions.mask_conn_comp}'}]},
                                    {'type': 'section',
                                     'id': 'chapter_3_3',
                                     'title': '3.3 相位堆栈',
                                     'children': [{'type': 'figure',
                                                   'asset': 'avg_phase_velocity',
                                                   'caption': '{templates.captions.avg_phase_velocity}',
                                                   'max_height_cm': 12}]},
                                    {'type': 'section',
                                     'id': 'chapter_3_4',
                                     'title': '3.4 相位解缠误差',
                                     'children': [{'type': 'figure',
                                                   'asset': 'unwrap_error',
                                                   'caption': '{templates.captions.unwrap_error}',
                                                   'max_height_cm': 12}]},
                                    {'type': 'section',
                                     'id': 'chapter_3_5',
                                     'title': '3.5 几何数据',
                                     'children': [{'type': 'figure',
                                                   'asset': 'geometry_radar',
                                                   'caption': '{templates.captions.geometry_radar}',
                                                   'max_height_cm': 12}]},
                                    {'type': 'section',
                                     'id': 'chapter_3_6',
                                     'title': '3.6 残余相位 RMS',
                                     'children': [{'type': 'figure',
                                                   'asset': 'residual_rms',
                                                   'caption': '{templates.captions.residual_rms}',
                                                   'max_height_cm': 12},
                                                  {'type': 'text_file_line',
                                                   'source': 'reference_dates',
                                                   'template': '{templates.text.reference_dates}'},
                                                  {'type': 'text_file_line',
                                                   'source': 'exclude_dates',
                                                   'template': '{templates.text.exclude_dates}',
                                                   'space_after_pt': 12}]}]},
                      {'type': 'chapter',
                       'id': 'chapter_4',
                       'title': '4. 形变结果',
                       'children': [{'type': 'section',
                                     'id': 'chapter_4_1',
                                     'title': '4.1 形变速率图',
                                     'children': [{'type': 'conditional',
                                                   'when': {'asset_exists': 'velocity_map'},
                                                   'then': [{'type': 'figure',
                                                             'asset': 'velocity_map',
                                                             'caption': '{templates.captions.velocity_map}'}],
                                                   'else': [{'type': 'paragraph',
                                                             'text': '[图片缺失: velocity_DEM.png 或 velocity.png]'}]}]},
                                    {'type': 'section',
                                     'id': 'chapter_4_2',
                                     'title': '4.2 VelocityStd',
                                     'children': [{'type': 'conditional',
                                                   'when': {'asset_exists': 'velocity_std'},
                                                   'then': [{'type': 'figure',
                                                             'asset': 'velocity_std',
                                                             'caption': '{templates.captions.velocity_std}'}],
                                                   'else': [{'type': 'paragraph',
                                                             'text': '[图片缺失: velocity_std_plot.png]'}]}]},
                                    {'type': 'section',
                                     'id': 'chapter_4_3',
                                     'title': '4.3 感兴趣目标点的时序形变图',
                                     'children': [{'type': 'conditional',
                                                   'when': {'source_nonempty': 'poi_figures'},
                                                   'then': [{'type': 'repeat',
                                                             'source': 'poi_figures',
                                                             'as': 'item',
                                                             'children': [{'type': 'figure',
                                                                           'path': '{item.path}',
                                                                           'caption': '{templates.captions.poi_timeseries}'}]}],
                                                   'else': [{'type': 'paragraph', 'text': '（未检测到感兴趣目标点图件）'}]}]}]},
                      {'type': 'chapter',
                       'id': 'chapter_5',
                       'title': '5. 解译分析',
                       'children': [{'type': 'interactive_text',
                                     'source': 'analysis_text',
                                     'line_spacing': 1.5,
                                     'first_line_indent_pt': 24,
                                     'font_size_pt': 12,
                                     'font_name_cn': '仿宋'}]}]}

# ========== 英文语言覆盖模板 ==========
EN_TEMPLATES = {
    "templates": {
        "headings": {
            "cover_title": (
                "Time-series InSAR Surface Deformation Processing and Analysis Report\n"
                "for {user_inputs.region_name} Using {user_inputs.data_source} Data"
            ),
        },
        "captions": {
            "network": (
                "Figure {figure_no} Interferogram network. Each node represents a SAR acquisition date, "
                "and each connecting line represents an interferometric pair. The horizontal axis shows time, "
                "and the vertical axis shows the perpendicular baseline relative to the reference image "
                "(unit: meters). The color of each connecting line indicates the mean spatial coherence "
                "of the corresponding interferometric pair, ranging from 0.2 to 1.0."
            ),
            "coherence_history": (
                "Figure {figure_no} Statistics of coherence extremes. The horizontal axis shows SAR acquisition "
                "dates, and the vertical axis shows the mean spatial coherence coefficient ranging from 0 to 1. "
                "For each acquisition date, the bar chart displays the extreme coherence statistics of all "
                "interferometric pairs associated with that image."
            ),
            "coherence_matrix": (
                "Figure {figure_no} Coherence matrix. The horizontal and vertical axes represent "
                "the index numbers of SAR images sorted by acquisition time."
            ),
            "avg_spatial_coh": "Figure {figure_no} Geocoded average spatial coherence map.",
            "temporal_coh": "Figure {figure_no} Geocoded temporal coherence map.",
            "mask_conn_comp": "Figure {figure_no} Phase unwrapping connected component mask.",
            "avg_phase_velocity": "Figure {figure_no} Mean phase velocity map.",
            "unwrap_error": "Figure {figure_no} Phase unwrapping error assessment.",
            "geometry_radar": "Figure {figure_no} Radar geometry parameters and auxiliary data.",
            "residual_rms": "Figure {figure_no} Residual phase RMS visualization.",
            "velocity_map": "Figure {figure_no} Deformation velocity map.",
            "velocity_std": "Figure {figure_no} Velocity standard deviation.",
            "poi_timeseries": "Figure {figure_no} Time-series deformation of Point of Interest {item_index}.",
        },
        "text": {
            "reference_dates": "Reference dates: {value}",
            "exclude_dates": "Excluded dates: {value}",
        },
    },
    "chapter_titles": {
        "chapter_1": "1. Data Information",
        "chapter_2": "2. Processing Parameters",
        "chapter_3": "3. Interferometric Data Quality Assessment",
        "chapter_4": "4. Deformation Results",
        "chapter_5": "5. Interpretation Analysis",
    },
    "section_titles": {
        "chapter_3_1": "3.1 Interferogram Network",
        "chapter_3_2": "3.2 Coherence",
        "chapter_3_3": "3.3 Phase Stack",
        "chapter_3_4": "3.4 Phase Unwrapping Error",
        "chapter_3_5": "3.5 Geometry",
        "chapter_3_6": "3.6 Residual Phase RMS",
        "chapter_4_1": "4.1 Deformation Velocity Map",
        "chapter_4_2": "4.2 Velocity Standard Deviation",
        "chapter_4_3": "4.3 Time-Series Deformation of Points of Interest",
    },
    "metadata_fields": {
        "satellite": "Satellite",
        "band": "Band",
        "wavelength": "Wavelength",
        "mode": "Acquisition Mode",
        "swath_Num": "Swath Number",
        "resolution": "Spatial Resolution",
        "track": "Track Number",
        "direction": "Orbit Direction",
        "inc_angle": "Incidence Angle",
        "az_angle": "Azimuth Angle",
        "pol": "Polarization",
        "start_time": "Start Date",
        "end_time": "End Date",
        "num_ifgram": "Number of Interferograms",
        "bbox": "Bounding Box",
    },
    "block_labels": {
        "software_versions": "Processing software",
    },
    "kv_labels": {
        "ts_method": "Time-series estimation method",
    },
    "interactive_defaults": {
        "ts_method": "SBAS",
        "analysis_text": "(Please add interpretation analysis.)",
    },
    "table_titles": {
        "obs_dates": "Table {table_no} Acquisition Dates of SAR Data",
        "mintpy_cfg": "MintPy Configuration Parameters",
    },
    "interactive_prompts": {
        "ts_method": ">>> Please enter the time-series estimation method (default: SBAS): ",
        "analysis_text": ">>> Please enter the interpretation analysis: ",
    },
    "cover_defaults": {
        "region_name": "Beijing",
        "output_name_template": "{user_inputs.region_name}_InSAR_Report.docx",
    },
    "defaults": {
        "font_cn": "Times New Roman",
        "font_heading_cn": "Times New Roman",
    },
    "sar_metadata": {
        "resolution_range": "Range",
        "resolution_azimuth": "Azimuth",
        "direction_descending": "Descending",
        "direction_ascending": "Ascending",
        "unit_meter": "m",
        "unit_degree": "deg",
    },
    "messages": {
        "software_not_installed": "Not Installed",
        "toc_placeholder": (
            "(Table of Contents: Right-click here -> Update Field -> "
            "Update Entire Table of Contents after opening in Word)"
        ),
        "toc_title": "Table of Contents",
        "cover_prompt_title": "Please enter the cover information:",
        "cover_prompt_source": "1. Please enter the data source",
        "cover_prompt_region": "2. Please enter the region name",
        "cover_prompt_author": "3. Please enter the lead author",
        "cover_organization": "AIRCAS",
        "cover_unit_label": "Organization:",
        "cover_author_label": "Author:",
        "cover_date_label": "Date:",
        "cover_date_format": "%B %d, %Y",
        "cover_config_done": "Configuration complete:",
        "cover_config_source": "Data source",
        "cover_config_region": "Region",
        "cover_config_author": "Lead author",
        "success_doc_generated": "[Success] Document generated: {path}",
        "success_toc_tip": (
            "Please open the Word document, right-click on the Table of Contents "
            "-> Update Field -> Update entire table of contents"
        ),
        "no_table_data": "(No table data detected)",
        "no_nondefault_config": "(No non-default configuration detected)",
        "no_poi_detected": "(No Points of Interest detected)",
        "missing_image": "[Missing image: {name}]",
        "missing_asset": "[Missing asset]",
        "not_found_or_empty": "[Not found or empty]",
        "table_header_param": "Parameter",
        "table_header_value": "Value",
        "software_label": "Processing software",
    },
}


# 检查给定的语言代码是否为英文。
def _is_en_lang(language):
    """Check if the given language code is English."""
    return language and str(language).startswith("en")


# 从配置或上下文中提取语言信息。
def _get_language(cfg_or_ctx):
    """Extract language from config or context dict."""
    if isinstance(cfg_or_ctx, dict):
        return str(cfg_or_ctx.get("project", {}).get("language", "zh-CN"))
    return "zh-CN"


# 当语言设置为英文时，对配置应用英文语言覆盖。
def apply_language_overrides(cfg):
    """Apply English language overrides to config if language is English."""
    lang = _get_language(cfg)
    if not _is_en_lang(lang):
        return cfg

    en = EN_TEMPLATES

    # 1. Template overrides (captions, headings, text)
    for category in ("headings", "captions", "text"):
        en_cat = en.get("templates", {}).get(category, {})
        for key, value in en_cat.items():
            cfg.setdefault("templates", {}).setdefault(category, {})[key] = value

    # 2. Chapter / section titles
    for block_id, title in en.get("chapter_titles", {}).items():
        block = find_block_by_id(cfg.get("report_structure", []), block_id)
        if block is not None:
            block["title"] = title
    for block_id, title in en.get("section_titles", {}).items():
        block = find_block_by_id(cfg.get("report_structure", []), block_id)
        if block is not None:
            block["title"] = title

    # 3. Metadata field labels
    field_labels = en.get("metadata_fields", {})
    _apply_metadata_field_labels(cfg, field_labels)

    # 4. Table titles (date grid, cfg table)
    table_titles = en.get("table_titles", {})
    _apply_table_titles(cfg, table_titles)

    # 5. Specific block labels (software_versions / kv_line)
    _apply_block_labels(
        cfg,
        block_labels=en.get("block_labels", {}),
        kv_labels=en.get("kv_labels", {}),
    )

    # 6. Interactive prompts + defaults
    prompts = en.get("interactive_prompts", {})
    for source_name, prompt_text in prompts.items():
        source = cfg.get("data_sources", {}).get(source_name)
        if isinstance(source, dict):
            source["prompt"] = prompt_text
    interactive_defaults = en.get("interactive_defaults", {})
    for source_name, default_value in interactive_defaults.items():
        source = cfg.get("data_sources", {}).get(source_name)
        if isinstance(source, dict):
            source["default"] = default_value

    # 7. Cover defaults
    cover_defaults = en.get("cover_defaults", {})
    for block in cfg.get("report_structure", []):
        if isinstance(block, dict) and block.get("type") == "cover":
            if "data_source" in cover_defaults:
                block["default_data_source"] = cover_defaults["data_source"]
            if "region_name" in cover_defaults:
                block["default_region_name"] = cover_defaults["region_name"]
            if "lead_author" in cover_defaults:
                block["default_lead_author"] = cover_defaults["lead_author"]
            if "output_name_template" in cover_defaults:
                block.setdefault("params", {})["output_name_template"] = cover_defaults["output_name_template"]
            break

    # 8. Defaults (fonts etc.)
    defaults = en.get("defaults", {})
    if defaults:
        cfg.setdefault("defaults", {}).update(defaults)

    # 9. Paragraph texts (missing images, placeholders) in report_structure
    para_texts = {
        "[图片缺失: velocity_DEM.png 或 velocity.png]": "[Missing image: velocity_DEM.png or velocity.png]",
        "[图片缺失: velocity_std_plot.png]": "[Missing image: velocity_std_plot.png]",
        "（未检测到感兴趣目标点图件）": _en_msg("no_poi_detected"),
    }
    _override_paragraph_text(cfg, para_texts)

    # 10. Quickstart English caption overrides
    qs_en_captions = cfg.pop("_captions_override_en", {})
    if qs_en_captions:
        templates = cfg.setdefault("templates", {}).setdefault("captions", {})
        for key, value in qs_en_captions.items():
            if isinstance(value, str) and value.strip():
                templates[key] = value.strip()

    # 11. YAML-level language_overrides.en (user customization, highest priority)
    yaml_en = cfg.get("language_overrides", {}).get("en", {})
    if yaml_en:
        cfg = _apply_yaml_en_overrides(cfg, yaml_en)

    return cfg


# 通过匹配数据键更新元数据表字段的显示标签。
def _apply_metadata_field_labels(cfg, field_labels):
    """Update metadata table field display labels by matching data keys."""
    def _walk(blocks):
        if not isinstance(blocks, list):
            return
        for block in blocks:
            if isinstance(block, dict):
                if block.get("type") == "metadata_table":
                    fields = block.get("fields", [])
                    new_fields = []
                    for pair in fields:
                        if isinstance(pair, list) and len(pair) == 2:
                            key = pair[1]
                            label = field_labels.get(key, pair[0])
                            new_fields.append([label, key])
                        else:
                            new_fields.append(pair)
                    block["fields"] = new_fields
                _walk(block.get("children", []))
    _walk(cfg.get("report_structure", []))


# 通过匹配数据源名称更新表格块的标题。
def _apply_table_titles(cfg, table_titles):
    """Update table block titles by matching source name."""
    def _walk(blocks):
        if not isinstance(blocks, list):
            return
        for block in blocks:
            if isinstance(block, dict):
                if block.get("type") in ("table", "cfg_table"):
                    source = block.get("source")
                    if source and source in table_titles:
                        block["title"] = table_titles[source]
                _walk(block.get("children", []))
    _walk(cfg.get("report_structure", []))



# 在生成英文报告时覆盖特定块类型的显示标签。
def _apply_block_labels(cfg, block_labels=None, kv_labels=None):
    """Override labels for specific block types when generating English reports."""
    block_labels = block_labels or {}
    kv_labels = kv_labels or {}

    def _walk(blocks):
        if not isinstance(blocks, list):
            return
        for block in blocks:
            if not isinstance(block, dict):
                continue

            if block.get("type") == "software_versions":
                label = block_labels.get("software_versions")
                if isinstance(label, str) and label.strip():
                    block["label"] = label.strip()

            if block.get("type") == "kv_line":
                source_name = block.get("source")
                if source_name in kv_labels and isinstance(kv_labels[source_name], str):
                    block["label"] = kv_labels[source_name].strip()

            _walk(block.get("children", []))
            _walk(block.get("then", []))
            _walk(block.get("else", []))

    _walk(cfg.get("report_structure", []))


# 在 report_structure 中用英文替换中文段落文本。
def _override_paragraph_text(cfg, text_map):
    """Override Chinese paragraph text with English equivalents in report_structure."""
    def _walk(blocks):
        if not isinstance(blocks, list):
            return
        for block in blocks:
            if isinstance(block, dict):
                if block.get("type") == "paragraph":
                    text = block.get("text", "")
                    if text in text_map:
                        block["text"] = text_map[text]
                _walk(block.get("children", []))
                _walk(block.get("then", []))
                _walk(block.get("else", []))
    _walk(cfg.get("report_structure", []))



# 应用 YAML 中 language_overrides.en 的用户自定义英文覆盖。
def _apply_yaml_en_overrides(cfg, yaml_en):
    """Apply user-customized English overrides from YAML language_overrides.en."""
    # Templates
    for category in ("headings", "captions", "text"):
        en_cat = yaml_en.get("templates", {}).get(category, {})
        for key, value in en_cat.items():
            if isinstance(value, str):
                cfg.setdefault("templates", {}).setdefault(category, {})[key] = value

    # Chapter/section titles
    for key in ("chapter_titles", "section_titles"):
        titles = yaml_en.get("report_structure_overrides", {}).get(key, {})
        for block_id, title in titles.items():
            block = find_block_by_id(cfg.get("report_structure", []), block_id)
            if block is not None and isinstance(title, str):
                block["title"] = title

    # Metadata field labels
    yaml_fields = yaml_en.get("metadata_fields", {})
    if yaml_fields:
        _apply_metadata_field_labels(cfg, yaml_fields)

    # Block labels / kv labels
    _apply_block_labels(
        cfg,
        block_labels=yaml_en.get("block_labels", {}),
        kv_labels=yaml_en.get("kv_labels", {}),
    )

    # Interactive defaults
    for source_name, default_value in yaml_en.get("interactive_defaults", {}).items():
        source = cfg.get("data_sources", {}).get(source_name)
        if isinstance(source, dict):
            source["default"] = default_value

    # Cover defaults
    yaml_cover = yaml_en.get("cover_defaults", {})
    for block in cfg.get("report_structure", []):
        if isinstance(block, dict) and block.get("type") == "cover":
            if "data_source" in yaml_cover:
                block["default_data_source"] = yaml_cover["data_source"]
            if "region_name" in yaml_cover:
                block["default_region_name"] = yaml_cover["region_name"]
            if "lead_author" in yaml_cover:
                block["default_lead_author"] = yaml_cover["lead_author"]
            if "output_name_template" in yaml_cover:
                block.setdefault("params", {})["output_name_template"] = yaml_cover["output_name_template"]
            break

    return cfg


# Module-level state for current language (used by font rendering)
_CURRENT_FONT_CN = "仿宋"
_CURRENT_LANGUAGE = "zh-CN"


# 设置模块级别的语言状态，供 set_font 等低层函数使用。
def _set_report_language(language):
    """Set the module-level language state used by set_font and other low-level functions."""
    global _CURRENT_FONT_CN, _CURRENT_LANGUAGE
    _CURRENT_LANGUAGE = str(language) if language else "zh-CN"
    if _is_en_lang(_CURRENT_LANGUAGE):
        _CURRENT_FONT_CN = "Times New Roman"
    else:
        _CURRENT_FONT_CN = "仿宋"


# 返回当前的模块语言。
def _get_report_language():
    return _CURRENT_LANGUAGE


# 判断当前语言是否为英文。
def _is_english():
    return _is_en_lang(_CURRENT_LANGUAGE)


# 根据模块语言状态返回中文或英文文本的简单内联翻译。
def _(zh_text, en_text):
    """Simple inline translation: return Chinese or English text based on module language state."""
    return en_text if _is_english() else zh_text


# 从 EN_TEMPLATES 中按 key 获取英文消息字符串。
def _en_msg(key):
    """Get English message string from EN_TEMPLATES by key, or return key itself if not found."""
    return EN_TEMPLATES.get("messages", {}).get(key, key)


# 严格按路径约定推导所有成图路径，无回退。
def resolve_figure_paths(base_dir, pic_dir):
    """严格按 picture_auto_max10.ipynb 的路径约定推导所有成图路径，无回退。

    保持与原始 Notebook 完全一致的路径结构：
      base_dir         → {project_dir}/mintpy
      geo_velocity     → {base_dir}/geo/geo_velocity.h5
      velocity_file    → {base_dir}/velocity.h5        (mintpy_meta 数据源 / velocityStd)
      timeseries_file  → {base_dir}/geo/geo_timeseries_ERA5_demErr.h5
      temporal_coh_file → {base_dir}/geo/geo_temporalCoherence.h5
      dem_file          → {base_dir}/../DEM/elevation_cop.dem
      pic_dir           → {base_dir}/pic
    """
    return {
        "velocity_file": os.path.join(base_dir, "velocity.h5"),
        "geo_velocity_file": os.path.join(base_dir, "geo", "geo_velocity.h5"),
        "timeseries_file": os.path.join(base_dir, "geo", "geo_timeseries_ERA5_demErr.h5"),
        "temporal_coh_file": os.path.join(base_dir, "geo", "geo_temporalCoherence.h5"),
        "dem_file": os.path.join(base_dir, "..", "DEM", "elevation_cop.dem"),
        "pic_dir": pic_dir,
    }


_SCHEMA_CACHE = {}
_PLACEHOLDER_RE = re.compile(r"\{([^{}]+)\}")


# 深度合并两个字典，返回合并后的副本。
def deep_merge(base, override):
    result = copy.deepcopy(base)
    for key, value in (override or {}).items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result


# 获取默认配置文件的候选路径列表，优先 quickstart。
def get_default_config_candidates():
    if "__file__" in globals():
        module_dir = os.path.dirname(os.path.abspath(__file__))
    else:
        module_dir = os.getcwd()

    # 方案四：优先 quickstart，找不到再用高级 schema
    return [
        os.path.join(module_dir, "report_quickstart.yaml"),
        os.path.join(module_dir, "report_schema.yaml"),
        os.path.join(module_dir, "captions.yaml"),
        os.path.join(os.getcwd(), "report_quickstart.yaml"),
        os.path.join(os.getcwd(), "report_schema.yaml"),
        os.path.join(os.getcwd(), "captions.yaml"),
    ]


# 验证配置字典包含所有必要的顶层字段。
def validate_schema(cfg):
    required_keys = [
        "schema_version",
        "runtime_paths",
        "defaults",
        "data_sources",
        "assets",
        "templates",
        "report_structure",
    ]
    for key in required_keys:
        if key not in cfg:
            raise ValueError(f"YAML 缺少顶层字段: {key}")

    if not isinstance(cfg["report_structure"], list):
        raise ValueError("report_structure 必须为列表。")


# 加载 YAML 文件并返回字典。
def load_yaml_dict(path):
    with open(path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    if not isinstance(data, dict):
        raise ValueError("YAML 顶层结构必须为字典。")
    return data


# 判断配置是否为 quickstart 模式。
def is_quickstart_config(cfg):
    return (
        isinstance(cfg, dict)
        and (
            str(cfg.get("schema_version", "")).startswith("quickstart")
            or "paths" in cfg
            or "cover_defaults" in cfg
            or "report_options" in cfg
        )
    )


# 判断配置是否为完整 schema 模式。
def is_full_schema_config(cfg):
    return (
        isinstance(cfg, dict)
        and "report_structure" in cfg
        and "runtime_paths" in cfg
        and "data_sources" in cfg
        and "assets" in cfg
    )


# 设置交互式数据源的默认值。
def set_interactive_default(full_cfg, source_name, default_value):
    if default_value is None:
        return
    source_cfg = full_cfg.get("data_sources", {}).get(source_name)
    if isinstance(source_cfg, dict) and source_cfg.get("type") == "interactive":
        source_cfg["default"] = default_value


# 递归搜索 report_structure 中指定 ID 的块。
def find_block_by_id(blocks, block_id):
    if not isinstance(blocks, list):
        return None
    for block in blocks:
        if isinstance(block, dict):
            if block.get("id") == block_id:
                return block
            found = find_block_by_id(block.get("children", []), block_id)
            if found is not None:
                return found
    return None


# 应用标题覆盖字典到 report_structure 中对应块。
def apply_titles_override(full_cfg, titles_override):
    if not isinstance(titles_override, dict):
        return
    for block_id, new_title in titles_override.items():
        block = find_block_by_id(full_cfg.get("report_structure", []), block_id)
        if block is not None and isinstance(new_title, str) and new_title.strip():
            block["title"] = new_title.strip()


# 从 report_structure 中移除指定 ID 的章节。
def remove_chapter(full_cfg, chapter_id):
    report_structure = full_cfg.get("report_structure", [])
    full_cfg["report_structure"] = [
        block for block in report_structure
        if not (isinstance(block, dict) and block.get("type") == "chapter" and block.get("id") == chapter_id)
    ]


# 从指定章节中移除指定 ID 的小节。
def remove_section(full_cfg, chapter_id, section_id):
    chapter = find_block_by_id(full_cfg.get("report_structure", []), chapter_id)
    if not chapter:
        return
    children = chapter.get("children", [])
    chapter["children"] = [
        block for block in children
        if not (isinstance(block, dict) and block.get("id") == section_id)
    ]


# 从 quickstart 配置加载基础完整 schema 模板。
def load_base_full_schema_from_quickstart(quickstart_path=None):
    # quickstart 与高级 schema 共存：如果 quickstart 同目录下有 report_schema.yaml，则优先用它作高级模板；
    # 否则退回 utils.py 内置模板。
    candidates = []
    if quickstart_path:
        quick_dir = os.path.dirname(os.path.abspath(quickstart_path))
        candidates.append(os.path.join(quick_dir, "report_schema.yaml"))
    if "__file__" in globals():
        candidates.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), "report_schema.yaml"))
    candidates.append(os.path.join(os.getcwd(), "report_schema.yaml"))

    for path in candidates:
        if path and os.path.exists(path):
            try:
                data = load_yaml_dict(path)
                if is_full_schema_config(data):
                    return deep_merge(FULL_SCHEMA_TEMPLATE, data), path
            except Exception as e:
                print(f"[警告] 读取高级 schema 失败，将回退到内置模板: {e}")

    return copy.deepcopy(FULL_SCHEMA_TEMPLATE), None


# 从 quickstart 配置构建完整的报告配置。
def build_full_schema_from_quickstart(quick_cfg, quickstart_path=None):
    full_cfg, schema_source = load_base_full_schema_from_quickstart(quickstart_path)

    if schema_source:
        print(f"[提示] quickstart 模式使用高级模板: {schema_source}")
    else:
        print("[提示] quickstart 模式未检测到外部 report_schema.yaml，使用 utils.py 内置完整模板。")

    # project
    project_name = quick_cfg.get("project_name")
    if isinstance(project_name, str) and project_name.strip():
        full_cfg.setdefault("project", {})["name"] = project_name.strip()

    # Language (quickstart 支持 language 字段，后续 apply_language_overrides 会应用)
    q_lang = quick_cfg.get("language")
    if q_lang:
        full_cfg.setdefault("project", {})["language"] = str(q_lang)

    # 路径
    paths = quick_cfg.get("paths", {})
    runtime_paths = full_cfg.setdefault("runtime_paths", {})
    if isinstance(paths, dict):
        if paths.get("logo_path"):
            runtime_paths["logo_path"] = paths["logo_path"]
        if paths.get("base_dir"):
            runtime_paths["base_dir"] = paths["base_dir"]
        if paths.get("output_dir"):
            runtime_paths["output_dir"] = paths["output_dir"]

    # 重新绑定依赖 base_dir / pic_dir 的默认派生路径
    runtime_paths["pic_dir"] = "{runtime_paths.base_dir}/pic"
    runtime_paths["velocity_file"] = "{runtime_paths.base_dir}/velocity.h5"
    runtime_paths["cfg_file"] = "{runtime_paths.pic_dir}/smallbaselineApp.cfg"
    runtime_paths["rms_txt"] = "{runtime_paths.pic_dir}/rms_timeseriesResidual_ramp.txt"
    runtime_paths["reference_date_txt"] = "{runtime_paths.pic_dir}/reference_date.txt"
    runtime_paths["exclude_date_txt"] = "{runtime_paths.pic_dir}/exclude_date.txt"

    # 封面默认值
    cover_defaults = quick_cfg.get("cover_defaults", {})
    cover_block = None
    for block in full_cfg.get("report_structure", []):
        if isinstance(block, dict) and block.get("type") == "cover":
            cover_block = block
            break

    if cover_block:
        if isinstance(cover_defaults, dict):
            if "data_source" in cover_defaults:
                cover_block["default_data_source"] = cover_defaults["data_source"]
            if "region_name" in cover_defaults:
                cover_block["default_region_name"] = cover_defaults["region_name"]
            if "lead_author" in cover_defaults:
                cover_block["default_lead_author"] = cover_defaults["lead_author"]
            if "output_name_template" in cover_defaults:
                cover_block.setdefault("params", {})["output_name_template"] = cover_defaults["output_name_template"]

    # 处理参数与解译分析默认值
    processing_defaults = quick_cfg.get("processing_defaults", {})
    if isinstance(processing_defaults, dict):
        set_interactive_default(full_cfg, "ts_method", processing_defaults.get("ts_method"))
        set_interactive_default(full_cfg, "analysis_text", processing_defaults.get("analysis_text"))

    # 报告选项
    report_options = quick_cfg.get("report_options", {})
    if isinstance(report_options, dict):
        if "auto_page_break_after_chapter" in report_options:
            full_cfg.setdefault("defaults", {})["auto_page_break_after_chapter"] = bool(report_options["auto_page_break_after_chapter"])

        if report_options.get("include_chapter_3") is False:
            remove_chapter(full_cfg, "chapter_3")
        if report_options.get("include_chapter_4") is False:
            remove_chapter(full_cfg, "chapter_4")
        if report_options.get("include_chapter_5") is False:
            remove_chapter(full_cfg, "chapter_5")

        include_poi_section = str(report_options.get("include_poi_section", "auto")).lower()
        if include_poi_section in ("false", "0", "no", "n"):
            remove_section(full_cfg, "chapter_4", "chapter_4_3")

    # 图注覆盖
    captions_override = quick_cfg.get("captions_override", {})
    templates = full_cfg.setdefault("templates", {}).setdefault("captions", {})
    if isinstance(captions_override, dict):
        for key, value in captions_override.items():
            if isinstance(value, str) and value.strip():
                templates[key] = value.strip()

    # 英文图注覆盖：存入 cfg 供 apply_language_overrides 使用
    captions_override_en = quick_cfg.get("captions_override_en", {})
    if isinstance(captions_override_en, dict) and captions_override_en:
        full_cfg["_captions_override_en"] = captions_override_en

    # 标题覆盖
    apply_titles_override(full_cfg, quick_cfg.get("titles_override", {}))

    # 允许通过 quickstart 传递少量 defaults 覆盖
    defaults_override = quick_cfg.get("defaults_override", {})
    if isinstance(defaults_override, dict):
        full_cfg["defaults"] = deep_merge(full_cfg.get("defaults", {}), defaults_override)

    # --- 新增：figure_generation 深度合并 ---
    q_fig = quick_cfg.get("figure_generation")
    if isinstance(q_fig, dict):
        full_fig = full_cfg.setdefault("figure_generation", {})
        if "mode" in q_fig:
            full_fig["mode"] = str(q_fig["mode"]).lower()
        if "manual_points" in q_fig:
            full_fig["manual_points"] = list(q_fig["manual_points"])
        if "auto_params" in q_fig and isinstance(q_fig["auto_params"], dict):
            full_fig.setdefault("auto_params", {}).update(q_fig["auto_params"])
        if "plot" in q_fig and isinstance(q_fig["plot"], dict):
            full_fig.setdefault("plot", {}).update(q_fig["plot"])

    # --- 新增：路径候选解析（不覆盖用户显式指定的值）---
    runtime_paths = full_cfg.setdefault("runtime_paths", {})
    bd = runtime_paths.get("base_dir", "")
    pd = runtime_paths.get("pic_dir", os.path.join(bd, "pic") if bd else "")
    if bd:
        derived = resolve_figure_paths(bd, pd)
        for key in ("velocity_file", "geo_velocity_file", "timeseries_file", "temporal_coh_file", "dem_file"):
            runtime_paths.setdefault(key, derived[key])

    full_cfg["_config_mode"] = "quickstart"
    full_cfg["_quickstart_path"] = quickstart_path
    full_cfg["_base_schema_source"] = schema_source
    return full_cfg


# 加载报告配置文件，支持 quickstart 和高级 schema 两种格式。
def load_report_config(yaml_path=None, force_reload=False):
    candidate_paths = [yaml_path] if yaml_path else get_default_config_candidates()

    resolved_yaml_path = None
    for path in candidate_paths:
        if path and os.path.exists(path):
            resolved_yaml_path = path
            break

    cache_key = resolved_yaml_path or "__default__"
    if not force_reload and cache_key in _SCHEMA_CACHE:
        return _SCHEMA_CACHE[cache_key]

    if resolved_yaml_path:
        raw_cfg = load_yaml_dict(resolved_yaml_path)
        if is_quickstart_config(raw_cfg):
            print(f"[提示] 已加载 quickstart 配置文件: {resolved_yaml_path}")
            cfg = build_full_schema_from_quickstart(raw_cfg, quickstart_path=resolved_yaml_path)
        elif is_full_schema_config(raw_cfg):
            print(f"[提示] 已加载高级 schema 配置文件: {resolved_yaml_path}")
            cfg = deep_merge(FULL_SCHEMA_TEMPLATE, raw_cfg)
            cfg["_config_mode"] = "advanced"
        else:
            raise ValueError("无法识别该 YAML 类型：既不是 quickstart，也不是完整 report_schema。")
    else:
        print("[提示] 未检测到配置文件，将使用 utils.py 内置完整模板。")
        cfg = copy.deepcopy(FULL_SCHEMA_TEMPLATE)
        cfg["_config_mode"] = "builtin"

    validate_schema(cfg)
    cfg["_yaml_path"] = resolved_yaml_path
    _SCHEMA_CACHE[cache_key] = cfg
    return cfg



# 按点分隔路径从嵌套字典中安全取值。
def get_by_dotted_path(data, path, default=None):
    if path is None:
        return default
    current = data
    for part in str(path).split("."):
        if isinstance(current, dict) and part in current:
            current = current[part]
        else:
            return default
    return current


# 解析字符串模板，替换其中的占位符。
def resolve_string_template(text, context):
    if not isinstance(text, str):
        return text

    def replacer(match):
        expr = match.group(1).strip()
        value = get_by_dotted_path(context, expr, default=match.group(0))
        if value is None:
            return ""
        return str(value)

    previous = None
    current = text
    max_rounds = 10
    rounds = 0
    while previous != current and rounds < max_rounds:
        previous = current
        current = _PLACEHOLDER_RE.sub(replacer, current)
        rounds += 1
    return current


# 递归解析对象中的所有字符串模板占位符。
def resolve_object(obj, context):
    if isinstance(obj, str):
        return resolve_string_template(obj, context)
    if isinstance(obj, list):
        return [resolve_object(x, context) for x in obj]
    if isinstance(obj, dict):
        return {k: resolve_object(v, context) for k, v in obj.items()}
    return obj


# 设置文档运行的字体的名称、大小、粗细和颜色。
def set_font(run, font_size_pt, is_bold=False, font_name_cn=None):
    if font_name_cn is None:
        font_name_cn = _CURRENT_FONT_CN
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name_cn)
    run.font.size = Pt(font_size_pt)
    run.font.bold = is_bold
    run.font.color.rgb = RGBColor(0, 0, 0)


# 为表格单元格设置底部边框线。
def set_cell_bottom_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    tcBorders.append(bottom)


# 在文档中创建 Word 目录域元素。
def create_toc_element(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    run._element.append(instrText)

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._element.append(fldChar_separate)

    run_display = paragraph.add_run()
    run_display.text = _(
        "（目录生成区：生成文档后，请右键点击此处 -> 更新域 -> 更新整个目录）",
        _en_msg("toc_placeholder"),
    )
    set_font(run_display, 10, is_bold=False)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_display._element.append(fldChar_end)


# 将 PDF 文件转换为 PNG 图片，缓存已存在的转换结果。
def convert_pdf_to_png(pdf_path, dpi=300):
    if not os.path.exists(pdf_path):
        print(f"[警告] 文件不存在: {pdf_path}")
        return None

    png_path = pdf_path.replace(".pdf", ".png")
    if os.path.exists(png_path) and os.path.getmtime(png_path) > os.path.getmtime(pdf_path):
        return png_path

    try:
        doc_pdf = fitz.open(pdf_path)
        page = doc_pdf.load_page(0)
        pix = page.get_pixmap(dpi=dpi)
        pix.save(png_path)
        doc_pdf.close()
        return png_path
    except Exception as e:
        print(f"[错误] PDF 转图片失败 ({pdf_path}): {e}")
        return None


# 安全读取文本文件，自动尝试多种编码。
def read_text_file_safely(*candidate_paths):
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb18030"]
    for path in candidate_paths:
        if not path or not os.path.exists(path):
            continue
        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    lines = [line.strip() for line in f.readlines() if line.strip()]
                return "; ".join(lines)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"[警告] 读取文件失败 ({path}): {e}")
                return None
    return None


# 从文本文件中读取日期列表。
def read_date_list_from_txt(file_path):
    date_list = []
    demo_text = """
    20230109    0.019
    20230121    0.022
    20230202    0.018
    20230214    0.017
    20230226    0.015
    """
    if os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
        except Exception as e:
            print(f"[警告] 读取 TXT 文件失败: {e}")
            return []
    else:
        print(f"[提示] 文件 {file_path} 不存在，使用模拟数据演示表格生成。")
        lines = demo_text.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        parts = line.split()
        if len(parts) >= 1:
            date_list.append(parts[0])
    return date_list


# 解析 SAR 元数据属性字典为格式化的元数据信息。
def parse_sar_metadata(atr):
    meta = {}
    lang = _get_report_language()
    is_en = _is_english()

    plat = atr.get("PLATFORM", "").lower()
    if "sen" in plat:
        meta["satellite"] = "Sentinel-1"
    elif "alos" in plat:
        meta["satellite"] = "ALOS"
    else:
        meta["satellite"] = atr.get("PLATFORM", "Unknown")

    wl = float(atr.get("WAVELENGTH", 0.056))
    meta["wavelength"] = f"{wl * 100:.1f} cm"
    if 0.024 <= wl <= 0.038:
        meta["band"] = "X"
    elif 0.038 <= wl <= 0.075:
        meta["band"] = "C"
    elif 0.150 <= wl <= 0.300:
        meta["band"] = "L"
    else:
        meta["band"] = "Unknown"

    meta["mode"] = atr.get("beam_mode", atr.get("ACQUISITION_MODE", "IW"))
    meta["swath_Num"] = atr.get("beam_swath", atr.get("beam_swath", "N/A"))

    rg_size = float(atr.get("rangePixelSize", atr.get("RANGE_PIXEL_SIZE", 0)))
    az_size = float(atr.get("azimuthPixelSize", atr.get("AZIMUTH_PIXEL_SIZE", 0)))
    if is_en:
        unit = "m" if rg_size > 1 else "deg"
        meta["resolution"] = f"{rg_size:.1f} {unit} (Range) × {az_size:.1f} {unit} (Azimuth)"
    else:
        unit = "米" if rg_size > 1 else "度"
        meta["resolution"] = f"{rg_size:.1f} {unit} (距离向) × {az_size:.1f} {unit} (方位向)"
    meta["track"] = atr.get("trackNumber", "N/A")

    orbit_dir = atr.get("ORBIT_DIRECTION", atr.get("passDirection", "")).upper()
    if "DESC" in orbit_dir:
        meta["direction"] = "Descending" if is_en else "降轨 (Descending)"
    elif "ASC" in orbit_dir:
        meta["direction"] = "Ascending" if is_en else "升轨 (Ascending)"
    else:
        heading = float(atr.get("HEADING", 0))
        if is_en:
            meta["direction"] = "Descending" if -180 <= heading <= -90 else "Ascending"
        else:
            meta["direction"] = "降轨 (Descending)" if -180 <= heading <= -90 else "升轨 (Ascending)"

    meta["inc_angle"] = f"{float(atr.get('CENTER_INCIDENCE_ANGLE', 0)):.2f}°"
    meta["az_angle"] = f"{float(atr.get('HEADING', 0)):.2f}°"
    meta["pol"] = atr.get("POLARIZATION", atr.get("polarization", "VV"))

    s_date = atr.get("START_DATE", "")
    e_date = atr.get("END_DATE", "")
    if len(s_date) == 8:
        s_date = f"{s_date[:4]}-{s_date[4:6]}-{s_date[6:8]}"
    if len(e_date) == 8:
        e_date = f"{e_date[:4]}-{e_date[4:6]}-{e_date[6:8]}"
    meta["start_time"] = s_date
    meta["end_time"] = e_date

    try:
        lats = [float(atr.get(f"LAT_REF{i}")) for i in range(1, 5)]
        lons = [float(atr.get(f"LON_REF{i}")) for i in range(1, 5)]
        meta["bbox"] = f"{min(lats):.2f}°N ~ {max(lats):.2f}°N, {min(lons):.2f}°E ~ {max(lons):.2f}°E"
    except Exception:
        meta["bbox"] = "N/A"

    meta["num_ifgram"] = atr.get("mintpy.networkInversion.numIfgram", "N/A")
    return meta


# 解析 MintPy 配置文件，返回参数键值对列表。
def parse_cfg_file(cfg_path):
    params = []
    if not os.path.exists(cfg_path):
        return params

    try:
        with open(cfg_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        for line in lines:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "=" in line:
                key, raw_val = line.split("=", 1)
                key = key.strip()
                val_clean = raw_val.split("#")[0].strip().strip("'").strip('"')
                if val_clean:
                    params.append((key, val_clean))
    except Exception as e:
        print(f"[错误] 解析配置文件 {os.path.basename(cfg_path)} 失败: {e}")
    return params


# 获取 ISCE 和 MintPy 软件的版本信息。
def get_software_versions():
    versions = {"ISCE": None, "MintPy": "Unknown"}

    try:
        import mintpy
        if hasattr(mintpy, "__version__"):
            versions["MintPy"] = mintpy.__version__
        elif hasattr(mintpy, "version"):
            versions["MintPy"] = mintpy.version.release_version
    except Exception:
        versions["MintPy"] = _("未安装", _en_msg("software_not_installed"))

    try:
        result = subprocess.run("isce_version.py", capture_output=True, text=True, shell=True)
        if result.returncode == 0 and result.stdout.strip():
            out = result.stdout.strip()
            versions["ISCE"] = out.split()[1] if len(out.split()) > 1 else out
    except Exception:
        pass

    if not versions["ISCE"]:
        versions["ISCE"] = "Unknown"
    return versions


# 在文档中添加一行文本段落。
def add_text_line(doc, text, font_size_pt=10.5, font_name_cn=None, space_after_pt=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    set_font(run, font_size_pt, is_bold=False, font_name_cn=font_name_cn)


# 在文档中插入图片及其图注，支持 PDF 自动转换。
def add_figure_with_caption(doc, image_path, caption_text, target_width_cm=14.0, max_height_cm=14.0, pdf_render_dpi=300):
    final_img_path = image_path
    if image_path.lower().endswith(".pdf"):
        converted_path = convert_pdf_to_png(image_path, dpi=pdf_render_dpi)
        if converted_path:
            final_img_path = converted_path
        else:
            err_text = _(
                f"[图片缺失: {os.path.basename(image_path)}]",
                _en_msg("missing_image").format(name=os.path.basename(image_path)),
            )
            p_err = doc.add_paragraph(err_text)
            p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p_err.runs:
                set_font(p_err.runs[0], 10.5, False, "Times New Roman")
            return

    if not os.path.exists(final_img_path):
        err_text = _(
            f"[图片缺失: {os.path.basename(final_img_path)}]",
            _en_msg("missing_image").format(name=os.path.basename(final_img_path)),
        )
        p_err = doc.add_paragraph(err_text)
        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_err.runs:
            set_font(p_err.runs[0], 10.5, False, "Times New Roman")
        return

    try:
        with PILImage.open(final_img_path) as img:
            w_px, h_px = img.size
            aspect_ratio = w_px / h_px if h_px else 1.0
            height_if_width_fixed = target_width_cm / aspect_ratio

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()

            if height_if_width_fixed > max_height_cm:
                run_img.add_picture(final_img_path, height=Cm(max_height_cm))
            else:
                run_img.add_picture(final_img_path, width=Cm(target_width_cm))
    except Exception as e:
        print(f"[警告] 图片尺寸计算失败，使用默认宽度插入: {e}")
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(final_img_path, width=Cm(target_width_cm))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(caption_text)
    set_font(run_cap, 10.5, is_bold=False, font_name_cn=None)


# 提示用户输入并返回输入值，为空时返回默认值。
def prompt_value(prompt, default=""):
    text = input(prompt).strip()
    return text or default


# 从文本中提取开头的数字序列。
def extract_leading_number(text, default="1"):
    match = re.match(r"^\s*(\d+)", str(text))
    return match.group(1) if match else str(default)


# 构建运行时上下文字典，包括路径解析和章节编号登记。
def build_runtime_context(cfg):
    ctx = {
        "config": cfg,
        "project": copy.deepcopy(cfg.get("project", {})),
        "runtime_paths": {},
        "defaults": copy.deepcopy(cfg.get("defaults", {})),
        "templates": copy.deepcopy(cfg.get("templates", {})),
        "data": {},
        "assets_resolved": {},
        "user_inputs": {},
        "output_name": None,
        "output_path": None,
        "chapter_order": [],
        "chapter_numbers": {},
        "counters": {"figure": {}, "table": {}},
    }

    runtime_paths = copy.deepcopy(cfg.get("runtime_paths", {}))
    # 多轮解析 runtime_paths 中的占位符
    for _ in range(10):
        prev = copy.deepcopy(runtime_paths)
        temp_context = copy.deepcopy(ctx)
        temp_context["runtime_paths"] = runtime_paths
        runtime_paths = resolve_object(runtime_paths, temp_context)
        if prev == runtime_paths:
            break
    ctx["runtime_paths"] = runtime_paths

    # 预先登记 chapter id -> chapter number
    chapter_index = 0
    for block in cfg.get("report_structure", []):
        if block.get("type") == "chapter":
            chapter_index += 1
            block_id = block.get("id", f"chapter_{chapter_index}")
            title = block.get("title", str(chapter_index))
            chapter_no = extract_leading_number(title, default=str(chapter_index))
            ctx["chapter_order"].append(block_id)
            ctx["chapter_numbers"][block_id] = chapter_no
    return ctx


# 根据数据源类型采集相应的数据。
def collect_data_source(source_name, source_cfg, ctx):
    source_cfg = resolve_object(copy.deepcopy(source_cfg), ctx)
    stype = source_cfg.get("type")

    if stype == "mintpy_attribute":
        from mintpy.utils import readfile
        atr = readfile.read_attribute(source_cfg["path"])
        return parse_sar_metadata(atr)

    if stype == "text_file":
        path = source_cfg["path"]
        parser = source_cfg.get("parser", "inline_text")
        if parser == "date_list":
            return read_date_list_from_txt(path)
        if parser == "inline_text":
            return read_text_file_safely(path)
        raise ValueError(f"未知 text_file parser: {parser}")

    if stype == "cfg_file":
        return parse_cfg_file(source_cfg["path"])

    if stype == "glob":
        pattern = source_cfg["pattern"]
        files = glob.glob(pattern)
        sort_mode = source_cfg.get("sort", "default")
        if sort_mode == "natural":
            files = sorted(files, key=natural_sort_key)
        else:
            files = sorted(files)
        return [{"path": p, "name": os.path.basename(p)} for p in files]

    if stype == "interactive":
        prompt = source_cfg.get("prompt", ">>> ")
        default = source_cfg.get("default", "")
        return prompt_value(prompt, default=default)

    if stype == "literal":
        return source_cfg.get("value")

    raise ValueError(f"未知数据源类型: {stype}")



# ========== 多页 / 编号图片自动收集工具 ==========
# MintPy/Matplotlib 在图件内容过多时，常会把同一类图拆成多张：
#   xxx_1.png, xxx_2.png, xxx_3.png, ...
# 本项目的 schema 通常只配置 xxx.png。下面的函数会在渲染图片时自动寻找这些编号变体，
# 并把它们作为同一 figure block 的多张图片依次插入 Word，图注保持完全一致。
_NUMBERED_SUFFIX_RE = re.compile(r"^(?P<stem>.+)_(?P<index>\d+)$")
_SUPPORTED_FIGURE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp", ".pdf"}


# 生成自然排序的键值，使 coherence_2 排在 coherence_10 之前。
def natural_sort_key(value):
    """Natural sort key: coherence_2.png comes before coherence_10.png."""
    return [int(x) if x.isdigit() else x.lower() for x in re.split(r"(\d+)", str(value))]


# 对路径列表去重并保持原有顺序。
def _dedupe_keep_order(items):
    seen = set()
    result = []
    for item in items:
        if not item:
            continue
        key = os.path.abspath(str(item))
        if key in seen:
            continue
        seen.add(key)
        result.append(str(item))
    return result


# 将文件名主干拆分为基础主干和编号后缀。
def _split_numbered_stem(stem):
    """
    Split a file stem into base stem and numeric suffix.

    Examples
    --------
    geo_velocity_2 -> (geo_velocity, 2)
    geo_velocity   -> (geo_velocity, None)
    """
    match = _NUMBERED_SUFFIX_RE.match(str(stem))
    if match:
        return match.group("stem"), int(match.group("index"))
    return str(stem), None


# 收集单个配置路径及其 _1/_2 等编号变体。
def _collect_numbered_variants_for_one_path(path, include_base=True):
    """
    Collect one configured figure path and its _1/_2/... variants.

    The function is intentionally conservative:
    - it searches only in the same directory;
    - it searches only files with the same extension as the configured path;
    - if the configured path is already xxx_1.png, it treats xxx as the base name
      and collects xxx_1.png, xxx_2.png, ...;
    - ordering is base image first, followed by numbered images in natural order.
    """
    if not path:
        return []

    path = str(path)
    directory = os.path.dirname(path) or "."
    filename = os.path.basename(path)
    stem, ext = os.path.splitext(filename)
    ext_lower = ext.lower()

    if ext_lower not in _SUPPORTED_FIGURE_EXTS:
        return [path] if os.path.exists(path) else []

    base_stem, numbered_index = _split_numbered_stem(stem)
    base_path = os.path.join(directory, base_stem + ext)
    pattern = os.path.join(directory, f"{base_stem}_[0-9]*{ext}")

    numbered_paths = []
    for candidate in glob.glob(pattern):
        c_stem, c_ext = os.path.splitext(os.path.basename(candidate))
        c_base_stem, c_index = _split_numbered_stem(c_stem)
        # glob 的 [0-9]* 也可能匹配奇怪名称，这里再严格确认一次。
        if c_ext.lower() == ext_lower and c_base_stem == base_stem and c_index is not None:
            numbered_paths.append(candidate)
    numbered_paths = sorted(numbered_paths, key=natural_sort_key)

    result = []
    if include_base and os.path.exists(base_path):
        result.append(base_path)
    result.extend(numbered_paths)

    # 如果用户显式配置的是 xxx_2.png，但由于某种原因上面的 pattern 没有收集到，至少保留该文件。
    if numbered_index is not None and os.path.exists(path):
        result.append(path)

    return _dedupe_keep_order(result)


# 返回应插入的所有图片路径，支持编号图片自动扩展。
def collect_figure_paths(path, expand_numbered=True):
    """
    Return all existing paths that should be inserted for a configured figure path.

    If expand_numbered=True, a configured xxx.png will also match xxx_1.png, xxx_2.png, ...
    This is the key compatibility layer for split MintPy figures.
    """
    if not path:
        return []
    if not expand_numbered:
        return [str(path)] if os.path.exists(str(path)) else []
    return _collect_numbered_variants_for_one_path(str(path), include_base=True)


# 从候选路径列表中解析实际存在的路径及其编号变体。
def collect_candidate_figure_paths(candidate_paths, expand_numbered=True):
    """
    Resolve a list of candidate paths using the existing fallback rule.

    Old behavior:
      choose the first existing candidate.

    New behavior:
      choose the first candidate whose base image OR numbered variants exist,
      then return all existing numbered variants belonging to that chosen candidate.
    """
    for path in candidate_paths or []:
        paths = collect_figure_paths(path, expand_numbered=expand_numbered)
        if paths:
            return paths, str(path)
    # Nothing exists. Return the first candidate as display/debug path for missing-message compatibility.
    fallback = str(candidate_paths[0]) if candidate_paths else None
    return [], fallback

# 采集配置中所有数据源的数据并存入上下文中。
def collect_all_data_sources(cfg, ctx):
    for name, source_cfg in cfg.get("data_sources", {}).items():
        ctx["data"][name] = collect_data_source(name, source_cfg, ctx)
    return ctx


# 解析资源配置，返回资源路径和存在状态等信息。
def resolve_asset(asset_name, ctx):
    if asset_name in ctx["assets_resolved"]:
        return ctx["assets_resolved"][asset_name]

    if asset_name not in ctx["config"].get("assets", {}):
        raise KeyError(f"assets 中未定义: {asset_name}")

    asset_cfg = resolve_object(copy.deepcopy(ctx["config"]["assets"][asset_name]), ctx)
    atype = asset_cfg.get("type")

    if atype == "figure":
        # 默认开启编号图片自动扩展：xxx.png -> xxx.png + xxx_1.png + xxx_2.png ...
        # 如果某个图确实只想插入单张，可在 report_schema.yaml 对应 asset 下设置：
        #   expand_numbered: false
        expand_numbered = bool(asset_cfg.get("expand_numbered", True))

        chosen_path = None
        chosen_paths = []
        if "path" in asset_cfg:
            chosen_path = asset_cfg["path"]
            chosen_paths = collect_figure_paths(chosen_path, expand_numbered=expand_numbered)
        elif "candidates" in asset_cfg:
            chosen_paths, chosen_path = collect_candidate_figure_paths(
                asset_cfg.get("candidates", []),
                expand_numbered=expand_numbered,
            )

        result = {
            "type": "figure",
            # path 保持旧接口兼容；paths 是新接口，用于插入多张编号图。
            "path": chosen_paths[0] if chosen_paths else chosen_path,
            "paths": chosen_paths,
            "exists": len(chosen_paths) > 0,
        }

    elif atype == "figure_group":
        source_name = asset_cfg.get("source")
        items = ctx["data"].get(source_name, [])
        result = {
            "type": "figure_group",
            "items": items,
            "exists": len(items) > 0,
        }
    else:
        raise ValueError(f"未知 asset 类型: {atype}")

    ctx["assets_resolved"][asset_name] = result
    return result


# 检查指定资源是否存在。
def asset_exists(asset_name, ctx):
    return resolve_asset(asset_name, ctx).get("exists", False)


# 检查指定数据源是否为非空。
def source_nonempty(source_name, ctx):
    value = ctx["data"].get(source_name)
    if value is None:
        return False
    if isinstance(value, (list, tuple, dict, str)):
        return len(value) > 0
    return bool(value)


# 递增指定类型的计数器并返回当前值。
def next_counter(counter_type, chapter_id, ctx):
    store = ctx["counters"].setdefault(counter_type, {})
    key = chapter_id or "__global__"
    store.setdefault(key, 0)
    store[key] += 1
    return store[key]


# 根据章节 ID 生成图编号。
def make_figure_no(chapter_id, ctx):
    idx = next_counter("figure", chapter_id, ctx)
    chapter_no = ctx["chapter_numbers"].get(chapter_id, "1")
    return f"{chapter_no}-{idx}"


# 根据章节 ID 生成表编号。
def make_table_no(chapter_id, ctx):
    idx = next_counter("table", chapter_id, ctx)
    chapter_no = ctx["chapter_numbers"].get(chapter_id, "1")
    return f"{chapter_no}-{idx}"


# 构建渲染上下文，包含项目信息、路径、模板和循环变量等。
def build_render_context(ctx, current_chapter_id=None, loop_vars=None, extra=None):
    render_ctx = {
        "project": ctx.get("project", {}),
        "runtime_paths": ctx.get("runtime_paths", {}),
        "defaults": ctx.get("defaults", {}),
        "templates": ctx.get("templates", {}),
        "data": ctx.get("data", {}),
        "user_inputs": ctx.get("user_inputs", {}),
        "current_chapter_id": current_chapter_id,
        "current_chapter_no": ctx["chapter_numbers"].get(current_chapter_id, ""),
        "output_name": ctx.get("output_name"),
        "output_path": ctx.get("output_path"),
    }
    if loop_vars:
        render_ctx.update(loop_vars)
    if extra:
        render_ctx.update(extra)
    return render_ctx


# 遍历并渲染整个报告结构中的所有块。
def render_report(doc, cfg, ctx):
    for block in cfg.get("report_structure", []):
        render_block(doc, block, ctx, current_chapter_id=None, loop_vars=None)


# 根据块类型分发到对应的渲染函数。
def render_block(doc, block, ctx, current_chapter_id=None, loop_vars=None):
    render_ctx = build_render_context(ctx, current_chapter_id=current_chapter_id, loop_vars=loop_vars)
    block = resolve_object(copy.deepcopy(block), render_ctx)
    btype = block.get("type")

    if btype == "cover":
        return render_cover(doc, block, ctx)

    if btype == "toc":
        return render_toc(doc, block, ctx)

    if btype in ("chapter", "section", "subsection"):
        return render_heading_block(doc, block, ctx, current_chapter_id)

    if btype == "paragraph":
        return render_paragraph_block(doc, block, ctx)

    if btype == "software_versions":
        return render_software_versions_block(doc, block, ctx)

    if btype == "kv_line":
        return render_kv_line_block(doc, block, ctx)

    if btype == "figure":
        return render_figure_block(doc, block, ctx, current_chapter_id, loop_vars)

    if btype == "metadata_table":
        return render_metadata_table_block(doc, block, ctx)

    if btype == "table":
        return render_table_block(doc, block, ctx, current_chapter_id)

    if btype == "cfg_table":
        return render_cfg_table_block(doc, block, ctx)

    if btype == "text_file_line":
        return render_text_file_line_block(doc, block, ctx)

    if btype == "interactive_text":
        return render_interactive_text_block(doc, block, ctx)

    if btype == "conditional":
        return render_conditional_block(doc, block, ctx, current_chapter_id, loop_vars)

    if btype == "repeat":
        return render_repeat_block(doc, block, ctx, current_chapter_id)

    if btype == "page_break":
        doc.add_page_break()
        return None

    raise ValueError(f"未知 block 类型: {btype}")


# 渲染报告封面，包含 Logo、标题和作者信息表格。
def render_cover(doc, block, ctx):
    params = copy.deepcopy(block.get("params", {}))
    title_template = ctx.get("templates", {}).get("headings", {}).get(
        "cover_title",
        "基于{data_source}数据的{region_name}地区\n时序InSAR地表形变处理与分析报告",
    )

    is_en = _is_english()
    cover_values = {}
    if is_en:
        prompts = {
            "data_source": _en_msg("cover_prompt_source"),
            "region_name": _en_msg("cover_prompt_region"),
            "lead_author": _en_msg("cover_prompt_author"),
        }
    else:
        prompts = {
            "data_source": "1. 请输入数据源",
            "region_name": "2. 请输入地区",
            "lead_author": "3. 请输入第一作者",
        }

    print(_en_msg("cover_prompt_title") if is_en else "请输入封面信息 ：")
    if is_en:
        print("------------------------------")
    for key, spec in params.items():
        if key in ("logo_path", "output_name_template"):
            continue
        if isinstance(spec, str) and spec == "interactive":
            default_key = f"default_{key}"
            default_val = block.get(default_key, "")
            prompt_label = prompts.get(key, f"Please enter {key}")
            default_label = "默认" if not is_en else "default"
            cover_values[key] = prompt_value(f"{prompt_label} [{default_label}: {default_val}]: ", default=default_val)
        else:
            cover_values[key] = spec

    ctx["user_inputs"].update(cover_values)

    output_name_template = params.get("output_name_template", "{region_name}InSAR处理报告.docx")
    output_name = resolve_string_template(output_name_template, build_render_context(ctx, extra=cover_values))
    ctx["output_name"] = output_name
    output_dir = ctx["runtime_paths"].get("output_dir", os.getcwd())
    ctx["output_path"] = os.path.join(output_dir, output_name)

    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = resolve_string_template(str(params.get("logo_path", "")), build_render_context(ctx))
    if logo_path and os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Cm(5))
    elif logo_path:
        print(f"[提示] 未找到 Logo 图片: {logo_path}")
    p_logo.paragraph_format.space_after = Pt(130)

    title_text = resolve_string_template(title_template, build_render_context(ctx, extra=cover_values))
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.5
    run_title = p_title.add_run(title_text)
    set_font(run_title, font_size_pt=24, is_bold=True, font_name_cn=None)
    p_title.paragraph_format.space_after = Pt(150)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False
    col_width_label = Cm(4.5)
    col_width_value = Cm(9.5)

    for row in table.rows:
        row.cells[0].width = col_width_label
        row.cells[1].width = col_width_value
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), "650")
        row._tr.get_or_add_trPr().append(tr_height)

    def fill_row(row_idx, label, value):
        cell_lbl = table.cell(row_idx, 0)
        p_l = cell_lbl.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_l = p_l.add_run(label)
        set_font(run_l, 18, is_bold=True, font_name_cn=None)
        cell_lbl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell_val = table.cell(row_idx, 1)
        p_v = cell_val.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_v = p_v.add_run(value)
        set_font(run_v, 18, is_bold=True, font_name_cn=None)
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if is_en:
        org = _en_msg("cover_organization")
        author_str = f"{cover_values.get('lead_author', 'yourname')}, YunJun Zhang"
        date_str = datetime.now().strftime(_en_msg("cover_date_format"))
        fill_row(0, _en_msg("cover_unit_label"), org)
        fill_row(1, _en_msg("cover_author_label"), author_str)
        fill_row(2, _en_msg("cover_date_label"), date_str)
    else:
        fill_row(0, "单　　位：", "中国科学院空天信息创新研究院")
        fill_row(1, "作　　者：", f"{cover_values.get('lead_author', 'yourname')}、张云俊")
        fill_row(2, "日　　期：", datetime.now().strftime("%Y 年 %m 月 %d 日"))

    print("-" * 30)
    if is_en:
        print(_en_msg("cover_config_done"))
        print(f"{_en_msg('cover_config_source')}: {cover_values.get('data_source', '')}")
        print(f"{_en_msg('cover_config_region')}: {cover_values.get('region_name', '')}")
        print(f"{_en_msg('cover_config_author')}: {cover_values.get('lead_author', '')}")
    else:
        print("参数设置完成：")
        print(f"数据源: {cover_values.get('data_source', '')}")
        print(f"地区: {cover_values.get('region_name', '')}")
        print(f"第一作者: {cover_values.get('lead_author', '')}")
    print("-" * 30)

    doc.add_page_break()


# 渲染报告目录页。
def render_toc(doc, block, ctx):
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.space_before = Pt(20)
    p_toc_title.paragraph_format.space_after = Pt(20)

    toc_title_text = _("目    录", _en_msg("toc_title"))
    run_toc_title = p_toc_title.add_run(toc_title_text)
    set_font(run_toc_title, font_size_pt=16, is_bold=True, font_name_cn=None)

    create_toc_element(doc)
    doc.add_page_break()


# 渲染章节或小节的标题块及其子元素。
def render_heading_block(doc, block, ctx, current_chapter_id=None):
    btype = block["type"]
    level = block.get("level")
    if level is None:
        level = 1 if btype == "chapter" else 2 if btype == "section" else 3

    title = block.get("title", "")
    heading = doc.add_heading(title, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run()
    font_size = 16 if level == 1 else 14 if level == 2 else 12
    heading_font = ctx.get("defaults", {}).get("font_heading_cn", "黑体")
    set_font(run, font_size, True, font_name_cn=heading_font)
    heading.paragraph_format.space_after = Pt(12 if level == 1 else 6)

    new_chapter_id = current_chapter_id
    if btype == "chapter":
        new_chapter_id = block.get("id", current_chapter_id)

    for child in block.get("children", []):
        render_block(doc, child, ctx, current_chapter_id=new_chapter_id)

    if btype == "chapter" and ctx["defaults"].get("auto_page_break_after_chapter", True):
        doc.add_page_break()


# 渲染段落文本块。
def render_paragraph_block(doc, block, ctx):
    text = block.get("text", "")
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = block.get("line_spacing", 1.5)
    if block.get("first_line_indent", False):
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.runs[0] if p.runs else p.add_run()
    set_font(run, block.get("font_size_pt", 12), block.get("is_bold", False), block.get("font_name_cn"))




# 渲染处理软件版本信息行。
def render_software_versions_block(doc, block, ctx):
    """
    渲染“处理软件”信息行。
    支持 schema 中的 block:
      - type: software_versions
        label: 处理软件
        platform_source: mintpy_meta
        stack_processor: auto
    """
    label = block.get("label", _("处理软件", _en_msg("software_label")))
    if _is_english() and label == "处理软件":
        label = _en_msg("software_label")
    platform_source = block.get("platform_source", "mintpy_meta")
    stack_processor = block.get("stack_processor", "auto")

    vers = get_software_versions()

    meta = ctx.get("data", {}).get(platform_source, {}) or {}

    if stack_processor == "auto":
        satellite = str(meta.get("satellite", "")).lower()
        if "sentinel" in satellite or "sen" in satellite:
            stack_proc = "topsStack"
        elif "alos" in satellite:
            stack_proc = "stripmapStack"
        else:
            stack_proc = "stripmapStack"
    else:
        stack_proc = str(stack_processor)

    soft_str = f"ISCE-2 ({vers['ISCE']}) / {stack_proc} + MintPy ({vers['MintPy']})"

    colon = ": " if _is_english() else "："

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5

    r1 = p.add_run(f"{label}{colon}")
    set_font(r1, 12, False)

    r2 = p.add_run(soft_str)
    set_font(r2, 12, False, 'Times New Roman')

    return None


# 渲染键值对信息行。
def render_kv_line_block(doc, block, ctx):
    label = block.get("label", "")
    source_name = block.get("source")
    value = ctx["data"].get(source_name, "") if source_name else block.get("value", "")

    if _is_english():
        if source_name == "ts_method" and label in ("", "时序估计方法"):
            label = "Time-series estimation method"

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    colon = ": " if _is_english() else "："
    r1 = p.add_run(f"{label}{colon}")
    set_font(r1, 12, False)
    r2 = p.add_run(str(value))
    set_font(r2, 12, False, "Times New Roman")


# 渲染元数据信息表格。
def render_metadata_table_block(doc, block, ctx):
    source_name = block.get("source")
    meta = ctx["data"].get(source_name, {})

    extra_pairs = block.get("extra_pairs", [])
    fields = list(block.get("fields", []))
    for pair in extra_pairs:
        if isinstance(pair, list) and len(pair) == 2:
            fields.append(pair)

    for label, key in fields:
        value = meta.get(key, "N/A")
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        colon = ": " if _is_english() else "："
        run_lbl = p.add_run(f"{label}{colon}")
        set_font(run_lbl, 12, False)
        run_val = p.add_run(str(value))
        set_font(run_val, 12, False, "Times New Roman")


# 渲染日期网格表格。
def render_date_grid_table(doc, title, values, columns, current_chapter_id, ctx):
    table_no = make_table_no(current_chapter_id, ctx)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)

    final_title = title
    if "{table_no}" in title:
        final_title = title.format(table_no=table_no)

    run_cap = p_cap.add_run(final_title)
    set_font(run_cap, 12, is_bold=False, font_name_cn=None)

    num_cols = max(1, int(columns))
    num_rows = max(1, math.ceil(len(values) / num_cols))
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = ctx["defaults"].get("table_style", "Table Grid")
    table.autofit = True

    for i, value in enumerate(values):
        row_idx = i // num_cols
        col_idx = i % num_cols
        cell = table.cell(row_idx, col_idx)
        cell.text = str(value)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.25
        for run in p.runs:
            set_font(run, 12, is_bold=False)


# 根据渲染器类型渲染表格块。
def render_table_block(doc, block, ctx, current_chapter_id):
    source_name = block.get("source")
    renderer = block.get("renderer", "date_grid")
    data = ctx["data"].get(source_name, [])
    if not data:
        p = doc.add_paragraph(_("（未检测到表格数据）", _en_msg("no_table_data")))
        run = p.runs[0] if p.runs else p.add_run()
        set_font(run, 10.5, False)
        return

    if renderer == "date_grid":
        render_date_grid_table(
            doc=doc,
            title=block.get("title", _("表{table_no} 数据表", "Table {table_no} Data Table")),
            values=data,
            columns=block.get("columns", 4),
            current_chapter_id=current_chapter_id,
            ctx=ctx,
        )
        return

    raise ValueError(f"未知 table renderer: {renderer}")


# 渲染配置参数表格。
def render_cfg_table_block(doc, block, ctx):
    title = block.get("title", _("参数配置", "Configuration Parameters"))
    source_name = block.get("source")
    cfg_params = ctx["data"].get(source_name, [])
    filter_cfg = block.get("filter", {})
    exclude_values = {str(v).lower() for v in filter_cfg.get("exclude_values", [])}

    filtered = []
    for k, v in cfg_params:
        if str(v).lower() in exclude_values:
            continue
        filtered.append((k, v))

    p_mint = doc.add_paragraph()
    p_mint.paragraph_format.space_before = Pt(12)
    p_mint.paragraph_format.space_after = Pt(6)
    r_mint = p_mint.add_run(title)
    set_font(r_mint, 12, True, font_name_cn="黑体")

    if not filtered:
        p_none = doc.add_paragraph(_("（未检测到非默认配置）", _en_msg("no_nondefault_config")))
        set_font(p_none.runs[0] if p_none.runs else p_none.add_run(), 10.5, False)
        return

    table = doc.add_table(rows=len(filtered) + 1, cols=2)
    table.style = ctx["defaults"].get("table_style", "Table Grid")
    table.autofit = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = _("参数 (Parameter)", _en_msg("table_header_param"))
    hdr_cells[1].text = _("值 (Value)", _en_msg("table_header_value"))
    for cell in hdr_cells:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.runs[0] if p.runs else p.add_run(), 12, True, "黑体")

    for i, (key, val) in enumerate(filtered):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(val)

        p_k = row_cells[0].paragraphs[0]
        p_k.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(p_k.runs[0] if p_k.runs else p_k.add_run(), 10.5, False, "Times New Roman")

        p_v = row_cells[1].paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p_v.runs[0] if p_v.runs else p_v.add_run(), 10.5, False, "Times New Roman")


# 渲染文本文件行块。
def render_text_file_line_block(doc, block, ctx):
    source_name = block.get("source")
    value = ctx["data"].get(source_name)
    if not value:
        value = _("[未找到或无内容]", _en_msg("not_found_or_empty"))
    template = block.get("template", "{value}")
    text = resolve_string_template(template, build_render_context(ctx, extra={"value": value}))
    add_text_line(
        doc,
        text,
        font_size_pt=block.get("font_size_pt", 10.5),
        font_name_cn=block.get("font_name_cn"),
        space_after_pt=block.get("space_after_pt", 6),
    )


# 渲染交互式文本输入块。
def render_interactive_text_block(doc, block, ctx):
    source_name = block.get("source")
    value = ctx["data"].get(source_name, "")
    if not value:
        value = block.get("default", "")
    p = doc.add_paragraph(value)
    p.paragraph_format.line_spacing = block.get("line_spacing", 1.5)
    p.paragraph_format.first_line_indent = Pt(block.get("first_line_indent_pt", 24))
    run = p.runs[0] if p.runs else p.add_run()
    set_font(run, block.get("font_size_pt", 12), False, block.get("font_name_cn"))


# 渲染图片块，支持编号图扩展和缺失策略。
def render_figure_block(doc, block, ctx, current_chapter_id, loop_vars=None):
    if "asset" in block:
        asset = resolve_asset(block["asset"], ctx)
        paths = list(asset.get("paths") or [])
        path_for_message = asset.get("path")
    else:
        path_for_message = block.get("path")
        # 对直接写 path 的 figure 也支持编号图扩展。
        # 为避免 repeat/glob 场景中 xxx_1.png、xxx_2.png 被重复展开，默认只对非编号文件名扩展；
        # 如果确实想对某个直接 path 强制扩展，可设置 expand_numbered: true。
        direct_expand = block.get("expand_numbered")
        if direct_expand is None:
            stem = os.path.splitext(os.path.basename(str(path_for_message or "")))[0]
            _, idx = _split_numbered_stem(stem)
            direct_expand = idx is None
        paths = collect_figure_paths(path_for_message, expand_numbered=bool(direct_expand))

    policy = block.get("on_missing", ctx["defaults"].get("missing_asset_policy", "warn"))
    if not paths:
        fname = os.path.basename(path_for_message) if path_for_message else "unknown"
        msg = _(
            f"[图片缺失: {fname}]",
            _en_msg("missing_image").format(name=fname),
        )
        if policy == "error":
            raise FileNotFoundError(msg)
        if policy == "warn":
            p = doc.add_paragraph(msg)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0] if p.runs else p.add_run(), 10.5, False, "Times New Roman")
        return

    # 一个 figure block 只生成一个 figure_no；该 block 展开的所有 _1/_2 编号图共用同一图注。
    figure_no = make_figure_no(current_chapter_id, ctx)
    render_ctx = build_render_context(ctx, current_chapter_id=current_chapter_id, loop_vars=loop_vars, extra={"figure_no": figure_no})
    caption = resolve_string_template(block.get("caption", ""), render_ctx)

    width_cm = float(block.get("width_cm", ctx["defaults"].get("figure_width_cm", 14.0)))
    max_height_cm = float(block.get("max_height_cm", ctx["defaults"].get("figure_max_height_cm", 14.0)))
    pdf_render_dpi = int(block.get("pdf_render_dpi", ctx["defaults"].get("pdf_render_dpi", 300)))

    for path in paths:
        add_figure_with_caption(
            doc,
            path,
            caption,
            target_width_cm=width_cm,
            max_height_cm=max_height_cm,
            pdf_render_dpi=pdf_render_dpi,
        )


# 根据条件渲染 then 或 else 分支。
def render_conditional_block(doc, block, ctx, current_chapter_id, loop_vars=None):
    when = block.get("when", {})
    ok = False
    if "asset_exists" in when:
        ok = asset_exists(when["asset_exists"], ctx)
    elif "source_nonempty" in when:
        ok = source_nonempty(when["source_nonempty"], ctx)
    else:
        raise ValueError("conditional.when 目前仅支持 asset_exists/source_nonempty")

    branch = block.get("then", []) if ok else block.get("else", [])
    for child in branch:
        render_block(doc, child, ctx, current_chapter_id=current_chapter_id, loop_vars=loop_vars)


# 遍历数据源列表循环渲染子块。
def render_repeat_block(doc, block, ctx, current_chapter_id):
    source_name = block.get("source")
    items = ctx["data"].get(source_name, [])
    alias = block.get("as", "item")
    for idx, item in enumerate(items, start=1):
        loop_vars = {
            alias: item,
            "item": item,
            "item_index": idx,
            "loop_index": idx,
        }
        for child in block.get("children", []):
            render_block(doc, child, ctx, current_chapter_id=current_chapter_id, loop_vars=loop_vars)



# 从 YAML 配置加载并构建完整的 Word 报告。
def build_report_from_yaml(doc, yaml_path=None, language=None):
    cfg = load_report_config(yaml_path=yaml_path)

    if language:
        cfg.setdefault("project", {})["language"] = language
    cfg.setdefault("project", {}).setdefault("language", "zh-CN")

    cfg = apply_language_overrides(cfg)
    _set_report_language(cfg.get("project", {}).get("language", "zh-CN"))

    ctx = build_runtime_context(cfg)
    ctx = collect_all_data_sources(cfg, ctx)
    render_report(doc, cfg, ctx)
    return doc, ctx


# ========== 成图函数（从 picture_auto_max10.ipynb / picture_manual.ipynb 迁移）==========

# 读取 HDF5 文件，返回数据和属性字典。
def read_h5_dataset(file_path, dataset_name=None):
    """读取 HDF5 文件，返回 (data, attrs)。dataset_name=None 时自动选择第一个 2D 数据集。"""
    import h5py
    with h5py.File(file_path, 'r') as f:
        if dataset_name is not None:
            if dataset_name in f:
                data = np.array(f[dataset_name])
                attrs = dict(f.attrs)
                return data, attrs
            raise KeyError(f'文件 {file_path} 中不存在数据集: {dataset_name}')
        for key in f.keys():
            arr = np.array(f[key])
            if arr.ndim == 2:
                attrs = dict(f.attrs)
                return arr, attrs
        raise ValueError(f'文件 {file_path} 中未找到二维数据集。')


# 将像素坐标转换为经纬度坐标。
def pixel_to_lalo(row, col, attrs):
    """将像素坐标 (row, col) 转为 (lat, lon)。"""
    lat = float(attrs['Y_FIRST']) + row * float(attrs['Y_STEP'])
    lon = float(attrs['X_FIRST']) + col * float(attrs['X_STEP'])
    return lat, lon


# 在掩码范围内对数组进行 min-max 归一化。
def normalize_array(arr, mask):
    """在 mask 范围内做 min-max 归一化。"""
    out = np.zeros_like(arr, dtype=np.float64)
    vals = arr[mask]
    if vals.size == 0:
        return out
    vmin = np.nanmin(vals)
    vmax = np.nanmax(vals)
    if np.isclose(vmax, vmin):
        out[mask] = 1.0
        return out
    out[mask] = (arr[mask] - vmin) / (vmax - vmin)
    return out


# 自动识别大形变区域并选取代表性目标点。
def auto_select_target_points(ctx, params=None):
    """自动识别大形变区域并选取代表点。返回 [(lat, lon), ...] 列表。"""
    import numpy as np
    from scipy.ndimage import label as ndimage_label

    if params is None:
        params = {}

    TEMPORAL_COH_THRESHOLD = params.get('temporal_coh_threshold', 0.70)
    MIN_REGION_PIXELS = params.get('min_region_pixels', 30)
    MAX_REGIONS = params.get('max_regions', 10)
    REGION_PERCENTILE_CANDIDATES = params.get('region_percentiles',
                                              [99.5, 99.0, 98.5, 98.0, 97.0, 95.0])

    rp = ctx['runtime_paths']
    velocity_file = rp.get('geo_velocity_file')
    temporal_coh_file = rp.get('temporal_coh_file')
    pic_dir = rp.get('pic_dir', '')

    if not velocity_file or not os.path.exists(velocity_file):
        print(f'[警告] velocity 文件不存在: {velocity_file}，跳过自动选点')
        return []
    if not temporal_coh_file or not os.path.exists(temporal_coh_file):
        print(f'[警告] temporal coherence 文件不存在: {temporal_coh_file}，跳过自动选点')
        return []

    print('[信息] 开始自动选点...')
    velocity, vel_attrs = read_h5_dataset(velocity_file, dataset_name='velocity')
    temporal_coh, _ = read_h5_dataset(temporal_coh_file, dataset_name='temporalCoherence')

    if velocity.shape != temporal_coh.shape:
        print(f'[警告] velocity 与 temporal coherence 尺寸不一致，跳过自动选点')
        print(f'       velocity file: {velocity_file}')
        print(f'       temporal_coh file: {temporal_coh_file}')
        print(f'       velocity shape: {velocity.shape}')
        print(f'       temporal_coh shape: {temporal_coh.shape}')
        return []

    abs_velocity = np.abs(velocity)
    valid_mask = (
        np.isfinite(abs_velocity) &
        np.isfinite(temporal_coh) &
        (temporal_coh >= TEMPORAL_COH_THRESHOLD)
    )

    if np.count_nonzero(valid_mask) == 0:
        print(f'[警告] 当前 temporal coherence 阈值 ({TEMPORAL_COH_THRESHOLD}) 下无有效像元')
        return []

    structure = np.ones((3, 3), dtype=np.int8)
    selected_labeled = None
    selected_threshold = None
    selected_num_regions = 0

    for pct in REGION_PERCENTILE_CANDIDATES:
        threshold = np.percentile(abs_velocity[valid_mask], pct)
        hotspot_mask = valid_mask & (abs_velocity >= threshold)
        labeled, num = ndimage_label(hotspot_mask, structure=structure)

        kept_region_ids = []
        for region_id in range(1, num + 1):
            region_mask = (labeled == region_id)
            if np.count_nonzero(region_mask) >= MIN_REGION_PIXELS:
                kept_region_ids.append(region_id)

        if kept_region_ids:
            filtered_labeled = np.zeros_like(labeled, dtype=np.int32)
            for new_id, old_id in enumerate(kept_region_ids, start=1):
                filtered_labeled[labeled == old_id] = new_id
            selected_labeled = filtered_labeled
            selected_threshold = threshold
            selected_num_regions = len(kept_region_ids)
            print(f'[信息] 分位数阈值: {pct}%，对应阈值 {threshold:.6f}，识别到 {selected_num_regions} 个形变区域')
            break

    if selected_labeled is None or selected_num_regions == 0:
        print('[信息] 未识别到满足条件的大形变区域')
        return []

    # 综合评分
    vel_norm = normalize_array(abs_velocity, valid_mask)
    coh_norm = normalize_array(temporal_coh, valid_mask)
    point_score = vel_norm * coh_norm

    region_infos = []
    for region_id in range(1, selected_num_regions + 1):
        region_mask = (selected_labeled == region_id)
        region_score = point_score.copy()
        region_score[~region_mask] = -np.inf
        best_flat_idx = np.argmax(region_score)
        best_row, best_col = np.unravel_index(best_flat_idx, region_score.shape)
        best_lat, best_lon = pixel_to_lalo(best_row, best_col, vel_attrs)

        region_abs_vel = abs_velocity[region_mask]
        region_coh = temporal_coh[region_mask]
        region_infos.append({
            'region_id': region_id,
            'pixels': int(np.count_nonzero(region_mask)),
            'mean_abs_velocity': float(np.nanmean(region_abs_vel)),
            'max_abs_velocity': float(np.nanmax(region_abs_vel)),
            'mean_temporal_coh': float(np.nanmean(region_coh)),
            'best_lat': float(best_lat),
            'best_lon': float(best_lon),
            'best_abs_velocity': float(abs_velocity[best_row, best_col]),
            'best_temporal_coh': float(temporal_coh[best_row, best_col]),
        })

    region_infos.sort(key=lambda x: (x['max_abs_velocity'], x['mean_temporal_coh'], x['pixels']), reverse=True)

    if MAX_REGIONS is not None:
        region_infos = region_infos[:MAX_REGIONS]

    selected_points = []
    for rank, info in enumerate(region_infos, start=1):
        selected_points.append((info['best_lat'], info['best_lon']))
        print(f"  [区域 {rank}] 原ID={info['region_id']}, 像元={info['pixels']}, "
              f"max|vel|={info['max_abs_velocity']:.6f}, meanCoh={info['mean_temporal_coh']:.4f}, "
              f"选点(lat, lon)=({info['best_lat']:.6f}, {info['best_lon']:.6f})")

    # 输出汇总文本
    if pic_dir:
        os.makedirs(pic_dir, exist_ok=True)
        summary_file = os.path.join(pic_dir, 'auto_selected_points_summary.txt')
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write(f'保留区域数 = {len(region_infos)}\n\n')
            for rank, info in enumerate(region_infos, start=1):
                f.write(f"[区域 {rank}] lat={info['best_lat']:.6f}, lon={info['best_lon']:.6f}, "
                        f"max|vel|={info['max_abs_velocity']:.6f}\n")

    return selected_points


# 生成 Velocity 标准差图。
def generate_velocity_std_figure(ctx):
    """生成 Velocity 标准差图。返回输出路径或 None。"""
    from mintpy.cli.view import main as view_main

    rp = ctx['runtime_paths']
    velocity_file = rp.get('velocity_file')
    pic_dir = rp.get('pic_dir', '')

    if not velocity_file or not os.path.exists(velocity_file):
        print('[信息] velocity 文件不存在，跳过 velocity_std 图')
        return None

    out_file = os.path.join(pic_dir, 'velocity_std_plot.png')
    os.makedirs(pic_dir, exist_ok=True)

    cmd_args = [
        velocity_file, 'velocityStd',
        '--unit', 'cm',
        '--figtitle', 'Velocity Standard Deviation',
        '--fontsize', '12',
        '-o', out_file,
    ]

    try:
        view_main(cmd_args)
        if os.path.exists(out_file):
            print(f'[成图] velocity_std: {out_file}')
            return out_file
    except SystemExit:
        pass
    except Exception as e:
        print(f'[警告] 生成 velocity_std 图失败: {e}')

    return None


# 生成叠加 DEM 的形变速率图。
def generate_velocity_dem_figure(ctx, plot_params=None):
    """生成叠加 DEM 的形变速率图。返回输出路径或 None。"""
    from mintpy.cli.view import main as view_main

    if plot_params is None:
        plot_params = {}

    rp = ctx['runtime_paths']
    velocity_file = rp.get('geo_velocity_file')
    dem_file = rp.get('dem_file')
    pic_dir = rp.get('pic_dir', '')

    if not velocity_file or not os.path.exists(velocity_file):
        print('[信息] velocity 文件不存在，跳过 velocity_DEM 图')
        return None

    out_file = os.path.join(pic_dir, 'velocity_DEM.png')
    os.makedirs(pic_dir, exist_ok=True)

    wrap_range = plot_params.get('wrap_range', [-3, 3])

    cmd_args = [
        velocity_file, 'velocity',
        '-u', 'cm',
        '--wrap', '--wrap-range',
        str(wrap_range[0]), str(wrap_range[1]),
        '--shade-az', str(plot_params.get('shade_az', 315)),
        '--shade-alt', str(plot_params.get('shade_alt', 45)),
        '--shade-frac', str(plot_params.get('shade_frac', 0.8)),
        '--base-color', str(plot_params.get('base_color', 0.5)),
        '--shade-exag', str(plot_params.get('shade_exag', 0.3)),
        '--scalebar', '0.2', '0.8', '0.1',
        '--lalo-label', '--ylabel-rot', '90',
        '-o', out_file,
        '--nodisplay',
    ]

    sub_lat = plot_params.get('sub_lat')
    sub_lon = plot_params.get('sub_lon')
    if sub_lat and len(sub_lat) == 2:
        cmd_args.extend(['--sub-lat', str(sub_lat[0]), str(sub_lat[1])])
    if sub_lon and len(sub_lon) == 2:
        cmd_args.extend(['--sub-lon', str(sub_lon[0]), str(sub_lon[1])])
    if dem_file and os.path.exists(dem_file):
        cmd_args.extend(['--dem', dem_file])

    try:
        view_main(cmd_args)
        if os.path.exists(out_file):
            print(f'[成图] velocity_DEM: {out_file}')
            return out_file
    except SystemExit:
        pass
    except Exception as e:
        print(f'[警告] 生成 velocity_DEM 图失败: {e}')

    return None


# 为单个目标点生成位置图和时序图的拼接图。
def generate_one_figure(ctx, idx, target_lat, target_lon, plot_params=None):
    """为单个目标点生成左图（位置图）+ 右图（时序图）拼接图。返回输出路径或 None。"""
    from mintpy.cli.view import main as view_main
    from mintpy.cli.tsview import main as tsview_main
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg

    if plot_params is None:
        plot_params = {}

    rp = ctx['runtime_paths']
    velocity_file = rp.get('geo_velocity_file')
    dem_file = rp.get('dem_file')
    ts_file = rp.get('timeseries_file')
    pic_dir = rp.get('pic_dir', '')

    if not velocity_file or not os.path.exists(velocity_file):
        print(f'[信息] velocity 文件不存在，跳过第 {idx} 个点')
        return None

    os.makedirs(pic_dir, exist_ok=True)
    temp_map_file = os.path.join(pic_dir, f'temp_map_view_{idx}.png')
    temp_ts_file = os.path.join(pic_dir, f'temp_ts_view_{idx}.png')
    final_file = os.path.join(pic_dir, f'final_figure{idx}.png')

    for f in [temp_map_file, temp_ts_file, final_file]:
        if os.path.exists(f):
            os.remove(f)

    plt.close('all')

    # --- 左侧位置图 ---
    wrap_range = plot_params.get('wrap_range', [-3, 3])
    map_cmd = [
        velocity_file, 'velocity',
        '--pts-lalo', str(target_lat), str(target_lon),
        '--pts-marker', 'ro', '--pts-ms', '10',
        '-u', 'cm',
        '--wrap', '--wrap-range',
        str(wrap_range[0]), str(wrap_range[1]),
        '--shade-az', str(plot_params.get('shade_az', 315)),
        '--shade-alt', str(plot_params.get('shade_alt', 45)),
        '--shade-frac', str(plot_params.get('shade_frac', 0.8)),
        '--base-color', str(plot_params.get('base_color', 0.5)),
        '--shade-exag', str(plot_params.get('shade_exag', 0.3)),
        '--scalebar', '0.2', '0.8', '0.1',
        '--lalo-label', '--ylabel-rot', '90',
        '--title', '', '-o', temp_map_file, '--save', '--nodisplay',
    ]

    sub_lat = plot_params.get('sub_lat')
    sub_lon = plot_params.get('sub_lon')
    if sub_lat and len(sub_lat) == 2:
        map_cmd.extend(['--sub-lat', str(sub_lat[0]), str(sub_lat[1])])
    if sub_lon and len(sub_lon) == 2:
        map_cmd.extend(['--sub-lon', str(sub_lon[0]), str(sub_lon[1])])
    if dem_file and os.path.exists(dem_file):
        map_cmd.extend(['--dem', dem_file])

    try:
        view_main(map_cmd)
    except SystemExit:
        pass
    except Exception as e:
        print(f'[警告] 第 {idx} 个点位置图失败: {e}')

    if not os.path.exists(temp_map_file):
        print(f'[信息] 第 {idx} 个点左图未生成，跳过')
        plt.close('all')
        return None

    plt.close('all')

    # --- 右侧时序图 ---
    if ts_file and os.path.exists(ts_file):
        ts_cmd = [
            ts_file, '--lalo', str(target_lat), str(target_lon),
            '--unit', 'cm', '--poly', '1',
            '--marker', 'o', '--ms', '4',
            '--figsize', '8', '6',
            '--title', 'Displacement Time-series',
            '--nodisplay',
        ]
        try:
            plt.close('all')
            tsview_main(ts_cmd)
            fig_ts = plt.gcf()
            fig_ts.savefig(temp_ts_file, bbox_inches='tight', dpi=150)
            plt.close('all')
        except SystemExit:
            plt.close('all')
        except Exception as e:
            print(f'[警告] 第 {idx} 个点时序图失败: {e}')
            plt.close('all')

    # --- 拼接左右图 ---
    if os.path.exists(temp_map_file) and os.path.exists(temp_ts_file):
        try:
            fig, axes = plt.subplots(1, 2, figsize=(18, 7),
                                     gridspec_kw={'width_ratios': [1, 1.2]})
            img_left = mpimg.imread(temp_map_file)
            axes[0].imshow(img_left)
            axes[0].axis('off')
            axes[0].set_title(f'Location Map\n(Lat: {target_lat}, Lon: {target_lon})', fontsize=14)

            img_right = mpimg.imread(temp_ts_file)
            axes[1].imshow(img_right)
            axes[1].axis('off')
            axes[1].set_title('Time Series Deformation (Linear Fit)', fontsize=14)

            plt.tight_layout()
            plt.savefig(final_file, dpi=150, bbox_inches='tight')
            plt.close('all')
            print(f'[成图] final_figure{idx}: {final_file}')
            return final_file
        except Exception as e:
            print(f'[警告] 第 {idx} 个点拼接失败: {e}')
            plt.close('all')

    # 时序图不存在时，只输出位置图
    if os.path.exists(temp_map_file) and not os.path.exists(temp_ts_file):
        try:
            fig, ax = plt.subplots(1, 1, figsize=(10, 7))
            img_left = mpimg.imread(temp_map_file)
            ax.imshow(img_left)
            ax.axis('off')
            ax.set_title(f'Location Map\n(Lat: {target_lat}, Lon: {target_lon})', fontsize=14)
            plt.tight_layout()
            plt.savefig(final_file, dpi=150, bbox_inches='tight')
            plt.close('all')
            print(f'[成图] final_figure{idx} (仅位置图): {final_file}')
            return final_file
        except Exception as e:
            print(f'[警告] 第 {idx} 个点位置图插入失败: {e}')
            plt.close('all')

    return None


# 成图总入口，根据配置调度所有图件生成。
def generate_all_figures(ctx, cfg):
    """成图总入口。根据 cfg.figure_generation 调度所有图件生成。返回生成图件路径列表。"""
    fig_cfg = cfg.get('figure_generation', {})
    mode = fig_cfg.get('mode', 'auto')

    if mode == 'skip':
        print('[信息] figure_generation.mode = skip，跳过成图')
        return []

    # 确定目标点
    if mode == 'auto':
        auto_params = fig_cfg.get('auto_params', {})
        target_points = auto_select_target_points(ctx, auto_params)
        if not target_points:
            print('[信息] 自动选点未找到目标点，仅生成基础图件')
    elif mode == 'manual':
        target_points = fig_cfg.get('manual_points', [])
        if not target_points:
            print('[信息] manual_points 为空，仅生成基础图件')
    else:
        print(f'[警告] 未知 mode: {mode}，跳过成图')
        return []

    plot_params = fig_cfg.get('plot', {})
    generated = []

    r = generate_velocity_std_figure(ctx)
    if r:
        generated.append(r)

    r = generate_velocity_dem_figure(ctx, plot_params)
    if r:
        generated.append(r)

    for i, (lat, lon) in enumerate(target_points, start=1):
        r = generate_one_figure(ctx, i, lat, lon, plot_params)
        if r:
            generated.append(r)

    return generated


# 加载配置、生成图件、构建并保存 Word 报告。
def build_report_and_save(yaml_path=None, language=None, generate_figures=True):
    cfg = load_report_config(yaml_path=yaml_path)

    # Apply language: explicit parameter > YAML config > Chinese default
    if language:
        cfg.setdefault("project", {})["language"] = language
    cfg.setdefault("project", {}).setdefault("language", "zh-CN")

    # Apply English overrides if needed
    cfg = apply_language_overrides(cfg)

    # Initialize module-level language state for low-level functions
    _set_report_language(cfg.get("project", {}).get("language", "zh-CN"))

    ctx = build_runtime_context(cfg)

    # ─── 新增：成图步骤（在采集数据源之前执行）───
    if generate_figures:
        fig_cfg = cfg.get("figure_generation", {})
        mode = fig_cfg.get("mode", "auto")
        if mode != "skip":
            print("[信息] 开始成图...")
            generated = generate_all_figures(ctx, cfg)
            print(f"[信息] 成图完成，共生成 {len(generated)} 个图件")

    ctx = collect_all_data_sources(cfg, ctx)

    output_dir = ctx["runtime_paths"].get("output_dir", os.getcwd())
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    render_report(doc, cfg, ctx)

    if not ctx.get("output_name"):
        ctx["output_name"] = "InSAR_Report.docx"
    if not ctx.get("output_path"):
        ctx["output_path"] = os.path.join(output_dir, ctx["output_name"])

    doc.save(ctx["output_path"])
    lang = _get_report_language()
    is_en = _is_english()

    separator = "-" * 30
    if is_en:
        print(separator)
        print(_en_msg("success_doc_generated").format(path=ctx["output_path"]))
        print(_en_msg("success_toc_tip"))
        print(separator)
    else:
        print(separator)
        print(f"[成功] 文档已生成: {ctx['output_path']}")
        print("操作提示：请打开 Word 文档，在目录页【右键 -> 更新域 -> 更新整个目录】")
        print(separator)
    return ctx["output_path"], ctx


if __name__ == "__main__":
    build_report_and_save()
